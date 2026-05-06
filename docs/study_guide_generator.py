"""Generate the ReviewPyPerAPI line-by-line study guide as a .docx file.

Run:  python3 docs/study_guide_generator.py

Produces: docs/ReviewPyPerAPI_line_by_line.docx

Format mirrors the CSCI E-28 line-by-line study guides (smsh, tarc, ssp):
title page, Part 0 prerequisite concepts, then per-file walkthroughs with
every meaningful line explained in plain English. Adds an architecture
section with ASCII callgraph and request flowchart.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────── helpers ───────────────────────


def _set_mono_font(run, size=9):
    run.font.name = "Menlo"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Menlo")
    rFonts.set(qn("w:hAnsi"), "Menlo")
    rFonts.set(qn("w:cs"), "Menlo")
    run.font.size = Pt(size)


def _shade(paragraph, color="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


class Doc:
    def __init__(self):
        self.d = Document()
        # Default body font size 11pt; left-align body to keep readable line lengths
        styles = self.d.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(11)

    # ---- structural ----
    def title(self, text):
        p = self.d.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(28)
        run.font.bold = True

    def subtitle(self, text):
        p = self.d.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.italic = True

    def author_block(self, lines):
        for line in lines:
            p = self.d.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)

    def page_break(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def h1(self, text):
        p = self.d.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(20)
        run.font.bold = True

    def h2(self, text):
        p = self.d.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(15)
        run.font.bold = True

    def h3(self, text):
        p = self.d.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True

    # ---- body ----
    def p(self, text):
        para = self.d.add_paragraph()
        para.add_run(text)

    def code(self, text):
        for line in text.split("\n"):
            para = self.d.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(line if line else " ")
            _set_mono_font(run, 9)
            _shade(para, "F4F4F4")

    def code_with_explain(self, code_lines, explain):
        """Show a block of code, then a plain-English explanation paragraph."""
        self.code(code_lines)
        self.p(explain)

    def callout(self, title, body):
        p = self.d.add_paragraph()
        r = p.add_run(title + ".  ")
        r.bold = True
        p.add_run(body)
        _shade(p, "FFF8E1")

    def bullets(self, items):
        for item in items:
            self.d.add_paragraph(item, style="List Bullet")

    def numbered(self, items):
        for item in items:
            self.d.add_paragraph(item, style="List Number")

    def save(self, path):
        self.d.save(path)


# ─────────────────────── content ───────────────────────


def title_page(doc: Doc):
    doc.title("ReviewPyPerAPI")
    doc.subtitle("Source Code Study Guide — Every Line Explained")
    doc.p("")
    doc.author_block(
        [
            "Niels Pacheco-Barrios",
            "Rolston Lab  ·  May 2026",
            "",
            "A complete walkthrough of the ReviewPyPerAPI codebase: a hosted web app that "
            "wraps the ReviewPyper toolkit for automated systematic literature reviews. "
            "Every meaningful line of the core files is explained in plain English, alongside "
            "a refresher on the React, FastAPI, and Tailwind concepts the code relies on. "
            "Includes an architecture section with a function callgraph and request flowchart. "
            "Designed so that for any line a reviewer points to, you can explain (1) what it does "
            "mechanically, (2) why it is there, and (3) what would break if it were removed.",
        ]
    )
    doc.page_break()


def toc(doc: Doc):
    doc.h1("Contents")
    items = [
        "Part 0  —  Web stack concepts you need",
        "Part 1  —  System architecture, callgraph, request flowchart",
        "Part 2  —  Backend  ·  reviewpyper_api/",
        "    Part 2.1  —  main.py (entry point and middleware wiring)",
        "    Part 2.2  —  middleware/test_mode.py (synthetic-data short-circuit)",
        "    Part 2.3  —  routers/files.py (uploads, project state, the chunked stream fix)",
        "Part 3  —  Frontend  ·  reviewpyper_frontend/",
        "    Part 3.1  —  main.tsx and App.tsx (boot sequence, providers)",
        "    Part 3.2  —  hooks/useConfig.ts (runtime configuration)",
        "    Part 3.3  —  services/api.ts (the HTTP layer)",
        "    Part 3.4  —  hooks/useApi.ts (TanStack Query wrappers)",
        "    Part 3.5  —  hooks/useProjectState.ts (localStorage state)",
        "    Part 3.6  —  components/ui/FileUpload.tsx (drag-drop, size, progress)",
        "    Part 3.7  —  components/PipelineStepper.tsx (progress indicator)",
        "    Part 3.8  —  pages/SetupReview.tsx (project + API access)",
        "    Part 3.9  —  pages/TitleScreening.tsx (representative pipeline page)",
        "    Part 3.10 —  pages/PdfProcessing.tsx (OA retrieval mode)",
        "Part 4  —  Build infrastructure",
        "    Part 4.1  —  Dockerfile.prod",
        "    Part 4.2  —  nginx.conf",
        "    Part 4.3  —  vite.config.ts",
        "    Part 4.4  —  index.css (Tailwind v4 design tokens)",
        "Part 5  —  Defense  ·  questions to be ready for",
        "Part 6  —  Validation harness  ·  validation/",
    ]
    for item in items:
        doc.p(item)
    doc.page_break()


# ────── Part 0: concepts ──────

def part0(doc: Doc):
    doc.h1("Part 0 — Web stack concepts you need")
    doc.p(
        "This section covers the language and framework features that appear throughout "
        "the ReviewPyPerAPI codebase. Read it once; then when you see these patterns in the "
        "code below, they will already make sense."
    )

    doc.h2("Browser, server, and HTTP")
    doc.p(
        "The application is split into two processes that talk to each other over HTTP. "
        "The browser runs the React frontend (compiled to JavaScript). When the user clicks a "
        "button, JavaScript code in the browser builds an HTTP request — a small text message "
        "with a method (GET, POST, PUT, DELETE), a path (e.g. /api/files/upload/abc123), "
        "headers, and an optional body — and sends it across the network to the server. The "
        "server (FastAPI in our case) receives the request, runs Python code, and sends back an "
        "HTTP response (a status code like 200 or 404, headers, and a body, usually JSON)."
    )
    doc.p(
        "Every request is independent. The server keeps no \"connection\" with the browser "
        "between requests; if state needs to persist across requests, it must be stored "
        "somewhere — in the database, on disk, or in the browser's localStorage."
    )

    doc.h2("JSON")
    doc.p(
        "JSON (JavaScript Object Notation) is a text format for structured data. It looks like "
        "JavaScript object literals: { \"name\": \"Niels\", \"age\": 28 }. Both Python and "
        "JavaScript can convert JSON to and from native data structures with one function call "
        "(json.loads / json.dumps in Python; JSON.parse / JSON.stringify in JS). All API "
        "responses in ReviewPyPerAPI are JSON."
    )

    doc.h2("multipart/form-data and file uploads")
    doc.p(
        "Plain JSON bodies cannot carry binary file data efficiently. For file uploads we use "
        "the multipart/form-data Content-Type. The body is divided into parts, each separated by "
        "a unique boundary string the client picks. Each part has its own headers and body. The "
        "server parses the parts back into separate fields. The browser's FormData API "
        "constructs this for you; axios converts a FormData instance into a multipart body "
        "automatically — and crucially, axios picks the boundary string itself. If you set the "
        "Content-Type header by hand without the boundary, the body becomes unparseable. (This "
        "was the actual bug behind the original \"big CSV upload broken\" complaint; see "
        "Part 3.3.)"
    )

    doc.h2("CORS — Cross-Origin Resource Sharing")
    doc.p(
        "Browsers refuse to let JavaScript on origin A make requests to origin B by default; "
        "they enforce the Same-Origin Policy. To allow cross-origin requests, the SERVER must "
        "send specific headers (Access-Control-Allow-Origin, Access-Control-Allow-Methods, "
        "etc.) saying \"yes, I accept requests from origin A.\" FastAPI's CORSMiddleware adds "
        "these headers. In dev, frontend (origin localhost:5173) and API (localhost:8000) are "
        "different origins, so CORS is required."
    )

    doc.h2("REST and routers")
    doc.p(
        "A REST API exposes resources at predictable URLs and uses the HTTP methods to mean "
        "different things on the same URL. Example: GET /files/projects (list), POST "
        "/files/projects (create), DELETE /files/projects/{id} (remove a specific one). FastAPI "
        "calls the function that handles a URL+method combination a \"route handler.\" Routes "
        "are grouped into APIRouter objects (one per resource). The main app then "
        "include_router(...) for each one."
    )

    doc.h2("async / await")
    doc.p(
        "Both Python (asyncio + FastAPI) and JavaScript use async/await to express non-blocking "
        "I/O. An async function returns a Promise (JS) / Coroutine (Python). The await keyword "
        "pauses the function until the promised value is ready, but releases the runtime to do "
        "other work in the meantime. This is why it matters that file writes inside an async "
        "FastAPI handler use anyio.open_file (non-blocking) instead of plain open() — a "
        "blocking write inside an async function freezes the entire event loop, meaning no "
        "other request gets served until the write finishes."
    )

    doc.h2("React: components, JSX, props, state, hooks")
    doc.p(
        "A React component is just a function that returns JSX. JSX looks like HTML inside JS "
        "and compiles to function calls. Props are inputs to a component; state is data the "
        "component owns and re-renders when it changes. Hooks are functions starting with "
        "\"use\" that let you tap into React features from inside function components: "
        "useState (local state), useEffect (run code after render), useContext (read a "
        "context), useCallback / useMemo (cache references), useRef (a mutable container that "
        "survives re-renders without triggering one)."
    )
    doc.p(
        "The Rules of Hooks: you may only call hooks at the top level of a component (not in "
        "loops, conditions, or after early returns) and only from React functions. Violating "
        "this corrupts React's bookkeeping."
    )
    doc.p(
        "A common anti-pattern: copying derived data into state via useEffect. If you can "
        "compute a value from existing state or props, just compute it during render — "
        "do not add an effect that calls setState. The codebase originally had this in "
        "SetupReview.tsx; see Part 3.8 for the fix."
    )

    doc.h2("TypeScript in 30 seconds")
    doc.p(
        "TypeScript is JavaScript with static types. You annotate variables, parameters, and "
        "return values; the compiler (tsc) checks them. Common syntax: let n: number = 5; "
        "function f(s: string): boolean { ... }. An interface defines an object shape: "
        "interface Project { id: string; name: string }. A union type allows multiple "
        "options: type Provider = 'openai' | 'anthropic'. The ?? operator is "
        "\"nullish coalescing\" — a ?? b returns a unless a is null/undefined, in which case "
        "b. The ?. operator is \"optional chaining\" — obj?.foo returns undefined when obj is "
        "null/undefined instead of throwing."
    )

    doc.h2("axios — the HTTP client")
    doc.p(
        "axios is a JavaScript library that wraps fetch() with a friendlier API: automatic "
        "JSON parsing, request/response interceptors, upload progress callbacks, and "
        "automatic Content-Type detection for FormData bodies. The codebase uses one shared "
        "instance: const http = axios.create()."
    )

    doc.h2("TanStack Query (formerly React Query)")
    doc.p(
        "TanStack Query manages server state (data fetched from an API) inside React. The two "
        "primitives:"
    )
    doc.bullets(
        [
            "useQuery({ queryKey, queryFn }) — fetches data, caches it under queryKey, "
            "re-fetches on window focus or interval, returns { data, isLoading, error }.",
            "useMutation({ mutationFn }) — wraps an action that changes server state. "
            "Returns mutate() / mutateAsync() and { isPending, isError, error }.",
        ]
    )
    doc.p(
        "After a successful mutation, you usually invalidateQueries(['key']) so the next "
        "render of any component using that query re-fetches and shows fresh data."
    )

    doc.h2("FastAPI primitives")
    doc.p(
        "@app.get(\"/health\") above an async function makes it the handler for "
        "GET /health. Path parameters appear in the URL as {project_id} and arrive as a "
        "function parameter of the matching name. Form fields and uploaded files arrive via "
        "Form(...) and File(...) declarations in the function signature. Pydantic models "
        "(classes inheriting from BaseModel) define request/response schemas: FastAPI "
        "validates JSON input against them and serializes return values to JSON."
    )

    doc.h2("Tailwind CSS v4 and semantic tokens")
    doc.p(
        "Tailwind is a CSS framework where every styling decision becomes a utility class "
        "(bg-white, p-4, rounded-lg, flex, gap-2). Instead of writing CSS, you compose "
        "classes on the element. Tailwind v4 adds a CSS-first @theme block where you define "
        "design tokens (colors, fonts, radii) as CSS variables. The codebase uses semantic "
        "tokens (--color-foreground, --color-primary) so swapping themes only requires "
        "changing the @theme block."
    )

    doc.h2("Docker and docker-compose")
    doc.p(
        "Docker packages an application together with its OS-level dependencies into an "
        "image. A container is a running instance of an image. docker-compose is a file "
        "format (docker-compose.yml) that defines a multi-container application: which "
        "images to run, how they connect, what ports to expose, what volumes to mount."
    )
    doc.p(
        "ReviewPyPerAPI uses three compose variants: docker-compose.yml (just the API), "
        "docker-compose.dev.yml (full stack — frontend + API for local dev), and "
        "docker-compose.real.yml (E2E test variant with real OpenAI calls)."
    )

    doc.h2("Volumes and DATA_DIR")
    doc.p(
        "Containers have ephemeral filesystems — when a container restarts, anything written "
        "inside is gone. Persistent data must live in a volume mounted from the host. The "
        "API container declares a volume at /data (controlled by the DATA_DIR env var). All "
        "uploaded files, project state.json files, and pipeline output files live under "
        "DATA_DIR. This is the project's storage seam."
    )

    doc.h2("nginx as the production proxy")
    doc.p(
        "In production, the React app is built into static HTML/JS/CSS and served by nginx. "
        "nginx also reverse-proxies /api/* to the FastAPI container. Two nginx settings "
        "matter for big uploads: client_max_body_size (rejects requests above this size with "
        "413) and proxy_request_buffering (when off, nginx streams the body straight to "
        "FastAPI instead of buffering it to disk first). See Part 4.2."
    )

    doc.h2("Pipeline workflow — the seven steps")
    doc.p(
        "ReviewPyPer is a 7-step pipeline that turns a literature corpus into structured "
        "data. Each step's output is another step's input; paths are tracked in PipelineState:"
    )
    doc.numbered(
        [
            "Setup — define research question, pick provider; produces api_key_path.",
            "Title Screening — upload references CSV; AI rates each title; produces "
            "title_output_csv_path.",
            "Abstract Screening — upload abstracts; AI screens against criteria; produces "
            "screened_abstract_path and master_list_path.",
            "PDF Processing — download PDFs from PubMed (with optional Sci-Hub fallback), "
            "post-process, OCR; produces pdf_dir and ocr_output_dir.",
            "Text & Sections — preprocess raw text and label structural sections (methods, "
            "results, discussion); produces preprocessed_dir and json_dir.",
            "Inclusion Evaluation — score each paper against inclusion criteria; produces "
            "inclusion_automated_csv_path.",
            "Data Extraction — pull structured answers per included paper; produces "
            "extraction_automated_csv_path.",
        ]
    )
    doc.page_break()


# ────── Part 1: architecture ──────

def part1(doc: Doc):
    doc.h1("Part 1 — System architecture, callgraph, request flowchart")

    doc.h2("Topology")
    doc.p(
        "Two services, two directories. The frontend is a React + Vite single-page app "
        "served by nginx in production. The API is a FastAPI app embedding the ReviewPyper "
        "research toolkit (cloned from upstream Git at Docker build time). The frontend "
        "talks to the API only via /api/*; nginx (prod) or Vite (dev) proxies that prefix to "
        "the FastAPI container. All persistent files live under DATA_DIR (a Docker volume)."
    )

    doc.h2("Service layout (ASCII)")
    doc.code(
        """┌────────────────────────────────────────────────────────────────────┐
│  BROWSER                                                          │
│  ┌─────────────────────────────────────────────────────────────┐  │
│  │  React app  ·  reviewpyper_frontend/                        │  │
│  │  • SetupReview · TitleScreening · AbstractScreening · ...   │  │
│  │  • TanStack Query cache  ·  axios  ·  localStorage          │  │
│  └─────────────────────┬───────────────────────────────────────┘  │
└────────────────────────┼──────────────────────────────────────────┘
                         │ HTTPS  /api/*
                         ▼
┌────────────────────────────────────────────────────────────────────┐
│  PROD: nginx  (reviewpyper_frontend/Dockerfile + nginx.conf)      │
│   • serves built static files                                      │
│   • proxies /api/* → reviewpyper-api:8000                         │
│   • client_max_body_size 500m  ·  proxy_request_buffering off     │
└────────────────────────┬──────────────────────────────────────────┘
                         │
                         ▼
┌────────────────────────────────────────────────────────────────────┐
│  FastAPI  ·  reviewpyper_api/                                     │
│   main.py                                                          │
│    ├── CORSMiddleware   ──►  TestModeMiddleware                    │
│    ├── routers/files     (upload, list, delete, project state)    │
│    ├── routers/titles    ┐                                         │
│    ├── routers/abstracts │                                         │
│    ├── routers/pdfs      │   each one wraps                       │
│    ├── routers/text      ├──►  ReviewPyper / logic/*.py            │
│    ├── routers/sections  │                                         │
│    ├── routers/inclusion │                                         │
│    └── routers/extraction┘                                         │
└────────────────────────┬──────────────────────────────────────────┘
                         │ pip-installed Python module
                         ▼
┌────────────────────────────────────────────────────────────────────┐
│  ReviewPyper (upstream OSS toolkit, cloned at build time)         │
│  → calls OpenAI / Anthropic API as needed                          │
└────────────────────────────────────────────────────────────────────┘

   /data  (Docker volume, mounted into the API container)
   └── <project_id>/
       ├── api_key.txt     ── server-managed or user-supplied
       ├── state.json      ── PipelineState mirror (planned source of truth)
       ├── <uploaded CSVs, abstracts.txt, ...>
       └── output CSVs and JSON from each pipeline step
"""
    )

    doc.h2("Function callgraph — the upload happy path")
    doc.p(
        "The most-touched code path in the app is uploading a references CSV. Reading the "
        "callgraph top-to-bottom mirrors the runtime call stack."
    )
    doc.code(
        """USER drops a CSV onto <FileUpload>
  │
  ▼
FileUpload.handleDrop                      (components/ui/FileUpload.tsx)
  └─► handleFiles
        └─► validate                       (size check ≤ maxSizeMB)
        └─► props.onFiles
              │
              ▼
TitleScreening.handleUpload                (pages/TitleScreening.tsx)
  └─► useUploadFile.mutateAsync({...})
        │
        ▼
useUploadFile  (TanStack Query wrapper)    (hooks/useApi.ts)
  └─► filesApi.upload(config, projectId, file, '', onProgress)
        │
        ▼
filesApi.upload                            (services/api.ts)
  └─► axios.post(url, FormData, {onUploadProgress})
        │   axios picks the multipart boundary itself
        ▼
HTTP POST /api/files/upload/{projectId}    (network)
  │
  ▼ Vite dev proxy (or nginx in prod)
  │
  ▼
FastAPI route                              (routers/files.py)
  upload_file(project_id, file, subfolder)
    └─► _safe_path(project_id)             (path traversal guard)
    └─► anyio.open_file(dest, 'wb')
          └─► loop:  await file.read(1MB)  →  await f.write(chunk)
    └─► return UploadResponse(filename, path)
  │
  ▼ JSON response
TanStack Query updates the mutation cache
  │
  ▼
TitleScreening receives result
  └─► setUploadedPath(result.path)
  └─► updatePipelineState(projectId, { csv_path: result.path })
        │
        ▼
useProjectState  (writes to localStorage)  (hooks/useProjectState.ts)
"""
    )

    doc.h2("Request flowchart — title-screening run")
    doc.p(
        "Once a CSV is uploaded, hitting \"Run Title Screening\" follows this flow:"
    )
    doc.code(
        """┌─────────────────────────────────────────────────────────────────┐
│ User clicks \"Run Title Screening\"                                │
└──────────────────────┬──────────────────────────────────────────┘
                       │
                       ▼
   handleScreen()  (TitleScreening.tsx)
                       │
                       ▼
        useScreenTitles.mutateAsync({...params})
                       │
                       ▼
       validateOrThrow(ScreenTitlesParams, params)        zod schema
                       │
                  ┌────┴────┐
                  ▼         ▼
              valid       invalid → throws → toast.error(...)
                  │
                  ▼
        pipelineApi.screenTitles(config, params)
                       │
                       ▼
            POST /api/titles/screen   (axios)
                       │
                       ▼
   ┌───────────────────┴────────────────────┐
   │ TestModeMiddleware  (test=true?)        │
   │  yes ─► return synthetic CSV + JSON     │
   │  no  ─► continue                        │
   └───────────────────┬────────────────────┘
                       │
                       ▼
   routers/titles.py:screen()
                       │
                       ▼
   logic/titles.py     (calls ReviewPyper)
                       │
                       ▼
       OpenAI / Anthropic chat completions
                       │
                       ▼
       writes title_output_csv_path under /data/<project_id>/
                       │
                       ▼
   {output_csv_path: \"/data/.../titles.csv\"}
                       │
                       ▼ JSON
       setOutputPath(result.output_csv_path)
       updatePipelineState({title_output_csv_path: ...})
                       │
                       ▼
       PipelineStepper now shows step 2 done ✓
"""
    )
    doc.page_break()


# ────── Part 2: Backend ──────

def part2_intro(doc: Doc):
    doc.h1("Part 2 — Backend  ·  reviewpyper_api/")
    doc.p(
        "The backend is a single FastAPI app. main.py wires middleware and includes the eight "
        "routers. Each router owns one resource (files) or one pipeline step. Routes are thin: "
        "they parse the request, call into ReviewPyper / logic/ for the real work, and return a "
        "Pydantic response model. The TestModeMiddleware short-circuits pipeline routes when "
        "?test=true is on the URL — that lets the frontend's E2E suite run without spending "
        "real OpenAI tokens. We walk through the three core files in detail; the seven other "
        "routers (titles, abstracts, pdfs, text, sections, inclusion, extraction) all follow "
        "the same structural pattern as the file router and are inspected briefly at the end."
    )


def part2_main_py(doc: Doc):
    doc.h2("Part 2.1 — main.py")
    doc.p(
        "Entry point of the FastAPI app. Sets up Python's import path so the embedded "
        "ReviewPyper module is reachable, configures CORS, attaches the test-mode "
        "middleware, mounts every router, and exposes a /health endpoint."
    )

    doc.h3("Full listing")
    doc.code(
        """import os
import sys
from pathlib import Path

from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware

_api_root = Path(__file__).resolve().parent
if str(_api_root) not in sys.path:
    sys.path.insert(0, str(_api_root))

try:
    import ReviewPyper  # noqa: F401
except Exception as exc:
    raise RuntimeError(
        "ReviewPyper is not available. This API expects ReviewPyper to be installed "
        "via the Docker image build (git clone + PYTHONPATH)."
    ) from exc

from middleware.test_mode import TestModeMiddleware
from routers.files import router as files_router
from routers.titles import router as titles_router
from routers.abstracts import router as abstracts_router
from routers.pdfs import router as pdfs_router
from routers.text import router as text_router
from routers.sections import router as sections_router
from routers.inclusion import router as inclusion_router
from routers.extraction import router as extraction_router

app = FastAPI(title=\"ReviewPyper API\")

# CORS origins are env-driven so prod deployments can override the local defaults.
_default_origins = \"http://localhost:5173,http://localhost:3000\"
_cors_origins = [o.strip() for o in os.environ.get(\"CORS_ORIGINS\", _default_origins).split(\",\") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_credentials=True,
    allow_methods=[\"*\"],
    allow_headers=[\"*\"],
)
app.add_middleware(TestModeMiddleware)

app.include_router(files_router)
app.include_router(titles_router)
app.include_router(abstracts_router)
app.include_router(pdfs_router)
app.include_router(text_router)
app.include_router(sections_router)
app.include_router(inclusion_router)
app.include_router(extraction_router)


@app.get(\"/health\")
async def health():
    return {\"status\": \"ok\"}
"""
    )

    doc.h3("Imports and import-path setup")
    doc.code_with_explain(
        "import os\nimport sys\nfrom pathlib import Path",
        "Standard library only. os reads environment variables (CORS_ORIGINS later). sys is "
        "needed to mutate sys.path so we can import ReviewPyper from a non-standard location. "
        "Path is the modern, object-oriented filesystem-path API.",
    )
    doc.code_with_explain(
        "from fastapi import FastAPI\nfrom fastapi.middleware.cors import CORSMiddleware",
        "FastAPI is the application class. CORSMiddleware is the standard-issue CORS handler — "
        "we configure it to allow our frontend origins.",
    )
    doc.code_with_explain(
        "_api_root = Path(__file__).resolve().parent\n"
        "if str(_api_root) not in sys.path:\n"
        "    sys.path.insert(0, str(_api_root))",
        "Take the absolute path of THIS file (main.py) and grab its parent directory — that "
        "gives us the reviewpyper_api/ folder. If it isn't already in sys.path, prepend it. "
        "Why? When uvicorn boots us, the working directory might not be reviewpyper_api/, "
        "but our internal imports (from routers..., from middleware...) require it on the "
        "import path. This makes the file robust regardless of how it is launched.",
    )
    doc.code_with_explain(
        "try:\n"
        "    import ReviewPyper  # noqa: F401\n"
        "except Exception as exc:\n"
        "    raise RuntimeError(\n"
        "        \"ReviewPyper is not available. This API expects ReviewPyper to be installed \"\n"
        "        \"via the Docker image build (git clone + PYTHONPATH).\"\n"
        "    ) from exc",
        "ReviewPyper is the upstream OSS research toolkit. The Dockerfile.prod git-clones "
        "it into /app/ReviewPyper at build time and adds it to PYTHONPATH (see Part 4.1). "
        "This try/except probes whether that worked. If it didn't, we fail loudly at startup "
        "with a useful message instead of getting a confusing ModuleNotFoundError later when "
        "a route handler tries to call into it. The # noqa: F401 silences a linter warning "
        "about an \"unused\" import — the import IS the test.",
    )
    doc.code_with_explain(
        "from middleware.test_mode import TestModeMiddleware\n"
        "from routers.files import router as files_router\n"
        "from routers.titles import router as titles_router\n"
        "...",
        "Now that sys.path is right, pull in the local middleware and router modules. We "
        "rename router → files_router etc. so the include_router(...) calls below read "
        "naturally.",
    )

    doc.h3("App and middleware")
    doc.code_with_explain(
        "app = FastAPI(title=\"ReviewPyper API\")",
        "Create the FastAPI application instance. The title shows up in the auto-generated "
        "OpenAPI docs at /docs.",
    )
    doc.code_with_explain(
        "_default_origins = \"http://localhost:5173,http://localhost:3000\"\n"
        "_cors_origins = [o.strip() for o in os.environ.get(\"CORS_ORIGINS\", _default_origins).split(\",\") if o.strip()]",
        "Read CORS_ORIGINS from the environment, falling back to a comma-separated list of "
        "local dev origins. Split on commas, trim whitespace, drop empties. The list "
        "comprehension gives us a clean list[str] regardless of whether the env var has "
        "leading/trailing whitespace, blank entries, etc. Production deployments override "
        "this in their docker-compose with their real domain.",
    )
    doc.code_with_explain(
        "app.add_middleware(\n"
        "    CORSMiddleware,\n"
        "    allow_origins=_cors_origins,\n"
        "    allow_credentials=True,\n"
        "    allow_methods=[\"*\"],\n"
        "    allow_headers=[\"*\"],\n"
        ")",
        "Attach the CORS middleware. It will add Access-Control-Allow-* headers to "
        "responses going to a browser whose Origin matches the allow_origins list. "
        "allow_credentials=True is required if the frontend ever sends cookies. "
        "[\"*\"] for methods and headers means we accept everything; the security boundary "
        "is the origin allow-list itself.",
    )
    doc.code_with_explain(
        "app.add_middleware(TestModeMiddleware)",
        "Attach our custom middleware. Middlewares execute in reverse-add order on requests "
        "and forward order on responses, so this one runs FIRST on incoming requests and "
        "LAST on outgoing responses. (See Part 2.2 for what it actually does.)",
    )

    doc.h3("Router mounting")
    doc.code_with_explain(
        "app.include_router(files_router)\n"
        "app.include_router(titles_router)\n"
        "app.include_router(abstracts_router)\n"
        "app.include_router(pdfs_router)\n"
        "app.include_router(text_router)\n"
        "app.include_router(sections_router)\n"
        "app.include_router(inclusion_router)\n"
        "app.include_router(extraction_router)",
        "Each router owns a URL prefix (e.g. /files for files_router, declared inside "
        "routers/files.py as APIRouter(prefix=\"/files\")). After include, FastAPI knows about "
        "every route. Order doesn't matter for correctness here, but reads naturally as the "
        "pipeline order: files first (control plane), then the seven pipeline stages.",
    )

    doc.h3("/health")
    doc.code_with_explain(
        "@app.get(\"/health\")\nasync def health():\n    return {\"status\": \"ok\"}",
        "Tiny liveness endpoint. Used by the Docker compose healthcheck (in "
        "docker-compose.dev.yml) so the frontend container only starts after the API is "
        "responding. Also useful for load balancers and uptime monitors.",
    )
    doc.page_break()


def part2_test_mode(doc: Doc):
    doc.h2("Part 2.2 — middleware/test_mode.py")
    doc.p(
        "When ?test=true appears on a pipeline endpoint URL, this middleware intercepts the "
        "request, writes plausible synthetic output files to /data, and returns a JSON "
        "response — without invoking the real route handler (which would call the LLM and "
        "spend tokens). It exists so the Playwright E2E suite can exercise the full UI flow "
        "without paying for OpenAI / Anthropic. Critically, it creates real files on disk so "
        "the file router and PipelineState behave end-to-end exactly as in production."
    )

    doc.h3("Imports and constants")
    doc.code_with_explain(
        "from __future__ import annotations\n"
        "import csv\nimport json\nimport os\n"
        "from pathlib import Path\n"
        "from fastapi import Request\n"
        "from fastapi.responses import JSONResponse\n"
        "from starlette.middleware.base import BaseHTTPMiddleware\n",
        "from __future__ import annotations makes type hints lazy strings — useful when a "
        "module references types that aren't yet imported. csv writes the synthetic study "
        "CSV; json writes the labelled-sections JSON. BaseHTTPMiddleware is Starlette's "
        "(and thus FastAPI's) base class for request/response middleware.",
    )
    doc.code_with_explain(
        "DATA_DIR = Path(os.environ.get(\"DATA_DIR\", \"./data\")).resolve()",
        "Same DATA_DIR convention used everywhere. Resolved to an absolute path at module "
        "load so subsequent uses are stable regardless of cwd changes.",
    )

    doc.h3("Synthetic data")
    doc.code_with_explain(
        "SYNTHETIC_STUDIES = [\n  {\"PMID\": \"34567890\", \"Title\": \"...\", ...},\n  ...\n]",
        "A list of three plausible study records (DBS for OCD), with title-screening "
        "decisions (Include / Exclude) baked in. The middleware writes these into the "
        "appropriate CSV / JSON shape per pipeline step. They're hand-curated so reviewers "
        "demoing the app see realistic content, not lorem ipsum.",
    )

    doc.h3("Middleware class skeleton")
    doc.code_with_explain(
        "class TestModeMiddleware(BaseHTTPMiddleware):\n"
        "    async def dispatch(self, request: Request, call_next):\n"
        "        if request.query_params.get(\"test\") != \"true\":\n"
        "            return await call_next(request)\n"
        "        # else: handle in test mode → write synthetic files, return synthetic JSON",
        "BaseHTTPMiddleware requires you to override dispatch. It's called once per request. "
        "If the URL doesn't have ?test=true, fall through to call_next which delegates to "
        "the real route. Otherwise, branch on request.url.path to decide which synthetic "
        "response shape to build.",
    )
    doc.callout(
        "Why a middleware and not a route flag",
        "Putting the test-mode logic in the route handlers themselves would clutter the real "
        "code path with if-test-mode checks. A middleware keeps the production code clean "
        "and the test scaffolding cleanly separable.",
    )

    doc.h3("Path branching pattern")
    doc.code_with_explain(
        "if path.endswith(\"/titles/screen\"):\n"
        "    return _handle_titles(request, body)\n"
        "if path.endswith(\"/abstracts/screen\"):\n"
        "    return _handle_abstracts(request, body)\n"
        "...",
        "Each pipeline endpoint has a short helper that knows what synthetic output to "
        "produce. The helper writes a real CSV / JSON to <DATA_DIR>/<project_id>/, returns "
        "JSONResponse({...}) with the same shape the real handler would return.",
    )
    doc.p(
        "If you're walking through this in a meeting, the key idea to communicate is: real "
        "files are written to disk so the FILE router behaves identically (the upload page "
        "will list them, downloads will work). Only the LLM call is faked."
    )
    doc.page_break()


def part2_files_router(doc: Doc):
    doc.h2("Part 2.3 — routers/files.py")
    doc.p(
        "This is the most-touched backend file. It owns project lifecycle (create / list / "
        "delete), file upload + download + listing + delete, API-key persistence, and "
        "project-state read/write. The chunked async-streaming upload added recently is in "
        "this file."
    )

    doc.h3("Module docstring + imports")
    doc.code_with_explain(
        "\"\"\"File management routes.\n\n"
        "The browser cannot write directly to the shared /data volume.\n"
        "This router handles upload, download, listing, and deletion of files\n"
        "so the frontend can feed file paths into the ReviewPyPerAPI pipeline.\n"
        "\"\"\"",
        "The docstring states the responsibility in one sentence. The frontend deals with "
        "abstract file paths (strings) — it never sends file BYTES anywhere except through "
        "the upload endpoint here.",
    )
    doc.code_with_explain(
        "from __future__ import annotations\n\n"
        "import json\n"
        "import os\n"
        "import shutil\n"
        "import uuid\n"
        "from pathlib import Path\n"
        "from typing import Any\n\n"
        "import anyio\n"
        "from fastapi import APIRouter, HTTPException, UploadFile, File, Form\n"
        "from fastapi.responses import FileResponse\n"
        "from pydantic import BaseModel",
        "Standard library: json (project state), os (env vars), shutil (recursive delete), "
        "uuid (project ID generation), Path, typing.Any (state.json is dynamic). "
        "anyio is FastAPI's async runtime — open_file gives us non-blocking writes. "
        "From FastAPI: APIRouter (a sub-app), HTTPException (raises a clean 4xx/5xx), "
        "UploadFile / File / Form (request-body declarations), FileResponse (efficient file "
        "download). BaseModel is Pydantic's validated-data class.",
    )

    doc.h3("Router and constants")
    doc.code_with_explain(
        "router = APIRouter(prefix=\"/files\", tags=[\"files\"])\n\n"
        "DATA_DIR = Path(os.environ.get(\"DATA_DIR\", \"./data\")).resolve()\n\n"
        "# Stream uploads to disk in 1 MB chunks so a large CSV doesn't sit in RAM.\n"
        "UPLOAD_CHUNK_SIZE = 1024 * 1024",
        "APIRouter declares the URL prefix once; every @router.get / @router.post below is "
        "implicitly under /files. tags=[\"files\"] is purely for the OpenAPI docs grouping. "
        "The 1 MB chunk size is the sweet spot — small enough that one chunk fits in memory "
        "comfortably, large enough that we aren't doing a syscall per kilobyte.",
    )

    doc.h3("_safe_path — path-traversal guard")
    doc.code_with_explain(
        "def _safe_path(relative: str) -> Path:\n"
        "    \"\"\"Resolve a relative path under DATA_DIR, preventing traversal.\n"
        "\n"
        "    Uses Path.is_relative_to so '/data/foobar' is correctly rejected even\n"
        "    when DATA_DIR is '/data/foo' (string prefix matching has a known\n"
        "    false-positive on shared name prefixes).\n"
        "    \"\"\"\n"
        "    resolved = (DATA_DIR / relative).resolve()\n"
        "    if not resolved.is_relative_to(DATA_DIR):\n"
        "        raise HTTPException(status_code=400, detail=\"Invalid path\")\n"
        "    return resolved",
        "Critical security primitive. Any handler that takes a project_id or path argument "
        "from the URL passes it through here. The pipeline: join DATA_DIR with the user "
        "input, resolve symlinks and dot segments to a canonical absolute path, then verify "
        "the canonical path is still under DATA_DIR. If a user passes \"../../etc/passwd\" "
        "it would resolve to a path outside DATA_DIR and is_relative_to returns False, so "
        "we raise HTTP 400. The previous implementation used str.startswith which has a "
        "subtle bug: \"/data/foobar\".startswith(\"/data/foo\") is True even though "
        "/data/foobar is not inside /data/foo. is_relative_to does the right comparison.",
    )

    doc.h3("Pydantic response models")
    doc.code_with_explain(
        "class ProjectInfo(BaseModel):\n"
        "    project_id: str\n"
        "    path: str\n"
        "    api_key_path: str | None = None\n"
        "\n"
        "class FileInfo(BaseModel):\n"
        "    name: str\n"
        "    path: str\n"
        "    size: int\n"
        "    is_dir: bool\n"
        "\n"
        "class UploadResponse(BaseModel):\n"
        "    filename: str\n"
        "    path: str\n"
        "\n"
        "class ApiKeyResponse(BaseModel):\n"
        "    api_key_path: str",
        "Pydantic models do double duty: at runtime they validate input/output (raise if "
        "fields missing or wrong types); at OpenAPI-doc time they describe the schema for "
        "API consumers. We declare them once and use them as response_model=... on the "
        "decorator below — FastAPI will then serialize the dict-like return value through "
        "the model so the JSON output is always shape-correct.",
    )

    doc.h3("create_project")
    doc.code_with_explain(
        "@router.post(\"/projects\", response_model=ProjectInfo)\n"
        "async def create_project(name: str = Form(...)):\n"
        "    \"\"\"Create a new project folder in /data.\n"
        "\n"
        "    If the OPENAI_API_KEY env var is set, the key is automatically saved\n"
        "    into the project so users don't need to provide their own.\n"
        "    \"\"\"\n"
        "    project_id = uuid.uuid4().hex[:12]\n"
        "    project_dir = DATA_DIR / project_id\n"
        "    project_dir.mkdir(parents=True, exist_ok=True)\n"
        "    (project_dir / \".project_name\").write_text(name)\n"
        "\n"
        "    api_key_path = None\n"
        "    server_key = os.environ.get(\"OPENAI_API_KEY\", \"\").strip()\n"
        "    if server_key:\n"
        "        key_file = project_dir / \"api_key.txt\"\n"
        "        key_file.write_text(server_key)\n"
        "        api_key_path = str(key_file)\n"
        "\n"
        "    return ProjectInfo(project_id=project_id, path=str(project_dir), api_key_path=api_key_path)",
        "POST /files/projects with a name form field creates a new review project. "
        "Form(...) tells FastAPI to expect application/x-www-form-urlencoded — convenient "
        "because the frontend already builds FormData for everything else. We coin a 12-char "
        "hex id from a uuid (random + collision-resistant), make the directory, drop a "
        ".project_name marker file inside it (handy for debugging on disk). If the operator "
        "set OPENAI_API_KEY in the API container's env, also write that key into "
        "api_key.txt — this is the \"managed credits\" path (ADR-0005). Return the path so "
        "the frontend can store it in PipelineState.api_key_path. The whole function is "
        "idempotent only at the directory level; calling twice creates two projects.",
    )

    doc.h3("list_projects, delete_project")
    doc.code_with_explain(
        "@router.get(\"/projects\", response_model=list[ProjectInfo])\n"
        "async def list_projects():\n"
        "    \"\"\"List all project folders.\"\"\"\n"
        "    projects = []\n"
        "    if not DATA_DIR.exists():\n"
        "        return projects\n"
        "    for entry in sorted(DATA_DIR.iterdir()):\n"
        "        if entry.is_dir() and not entry.name.startswith(\".\"):\n"
        "            projects.append(ProjectInfo(project_id=entry.name, path=str(entry)))\n"
        "    return projects",
        "Walk DATA_DIR and emit one ProjectInfo per non-dotfile directory. sorted() gives "
        "deterministic order (alphabetical). The DATA_DIR.exists() guard handles fresh "
        "deployments where the volume is empty.",
    )
    doc.code_with_explain(
        "@router.delete(\"/projects/{project_id}\")\n"
        "async def delete_project(project_id: str):\n"
        "    project_dir = _safe_path(project_id)\n"
        "    if not project_dir.is_dir():\n"
        "        raise HTTPException(status_code=404, detail=\"Project not found\")\n"
        "    shutil.rmtree(project_dir)\n"
        "    return {\"deleted\": project_id}",
        "{project_id} in the URL becomes a function parameter of the same name. Pass it "
        "through _safe_path before doing anything destructive. shutil.rmtree recursively "
        "removes the directory and everything in it. We return a small JSON acknowledgement "
        "so the client knows the operation succeeded.",
    )

    doc.h3("upload_file — the chunked stream")
    doc.code_with_explain(
        "@router.post(\"/upload/{project_id}\", response_model=UploadResponse)\n"
        "async def upload_file(\n"
        "    project_id: str,\n"
        "    file: UploadFile = File(...),\n"
        "    subfolder: str = Form(\"\"),\n"
        "):\n"
        "    \"\"\"Upload a file into a project folder.\n"
        "\n"
        "    Streams the body to disk in chunks via anyio.to_thread so the event loop\n"
        "    stays responsive even for multi-hundred-MB CSV uploads.\n"
        "    \"\"\"\n"
        "    project_dir = _safe_path(project_id)\n"
        "    if not project_dir.is_dir():\n"
        "        raise HTTPException(status_code=404, detail=\"Project not found\")\n"
        "\n"
        "    if not file.filename:\n"
        "        raise HTTPException(status_code=400, detail=\"Missing filename\")\n"
        "\n"
        "    target_dir = project_dir / subfolder if subfolder else project_dir\n"
        "    target_dir.mkdir(parents=True, exist_ok=True)\n"
        "    dest = target_dir / file.filename\n"
        "\n"
        "    async with await anyio.open_file(dest, \"wb\") as f:\n"
        "        while chunk := await file.read(UPLOAD_CHUNK_SIZE):\n"
        "            await f.write(chunk)\n"
        "\n"
        "    return UploadResponse(filename=file.filename, path=str(dest))",
        "This is the function we rewrote. UploadFile is FastAPI's wrapper around a "
        "Starlette spooled tempfile — it gives us .filename and .read(n). The "
        "subfolder Form field is empty by default; pages can specify a subfolder if they "
        "want to organize uploads (none currently do). _safe_path again guards the path. "
        "The mkdir(parents=True, exist_ok=True) is idempotent — both flags matter: "
        "parents=True creates intermediate directories, exist_ok=True suppresses error if "
        "it already exists.",
    )
    doc.p(
        "The hot loop: anyio.open_file returns an awaitable async context manager. We "
        "await it (the outer await), then enter the async with (which itself implicitly "
        "calls __aenter__). Inside, we call file.read(1MB) — read is async, so we await "
        "it; it returns up to UPLOAD_CHUNK_SIZE bytes (or b\"\" at EOF). The walrus operator "
        "(:=) assigns chunk and tests truthiness in one expression — when chunk is empty "
        "bytes (falsy), we exit the loop. f.write(chunk) is also async (anyio writes don't "
        "block the event loop). The outer async with auto-closes the file when we leave."
    )
    doc.callout(
        "Why this matters",
        "The old implementation used shutil.copyfileobj which is synchronous. Inside an "
        "async route, sync I/O blocks the event loop — meaning while one user uploads a "
        "300 MB CSV, every other API request stalls. The async chunked stream solves that.",
    )

    doc.h3("download_file, list_files, delete_file")
    doc.code_with_explain(
        "@router.get(\"/download/{path:path}\")\n"
        "async def download_file(path: str):\n"
        "    \"\"\"Download a file from /data by its path.\"\"\"\n"
        "    file_path = _safe_path(path)\n"
        "    if not file_path.is_file():\n"
        "        raise HTTPException(status_code=404, detail=\"File not found\")\n"
        "    return FileResponse(file_path, filename=file_path.name)",
        "The {path:path} converter tells FastAPI to capture EVERYTHING after /download/ "
        "(including slashes) as one parameter — without :path it stops at the first slash. "
        "FileResponse handles the heavy lifting: opens the file, sets Content-Type, sets "
        "Content-Disposition: attachment; filename=..., streams the bytes.",
    )
    doc.code_with_explain(
        "@router.get(\"/list/{project_id}\", response_model=list[FileInfo])\n"
        "async def list_files(project_id: str, subfolder: str = \"\"):\n"
        "    target = _safe_path(f\"{project_id}/{subfolder}\" if subfolder else project_id)\n"
        "    if not target.is_dir():\n"
        "        raise HTTPException(status_code=404, detail=\"Directory not found\")\n"
        "    entries = []\n"
        "    for entry in sorted(target.iterdir()):\n"
        "        if entry.name.startswith(\".\"):\n"
        "            continue\n"
        "        entries.append(FileInfo(\n"
        "            name=entry.name,\n"
        "            path=str(entry),\n"
        "            size=entry.stat().st_size if entry.is_file() else 0,\n"
        "            is_dir=entry.is_dir(),\n"
        "        ))\n"
        "    return entries",
        "Lists files in a project folder, optionally narrowed to a subfolder via a query "
        "param (?subfolder=PDFs). Skip dotfiles (.project_name and friends). Size 0 is "
        "reported for directories — they don't have a meaningful single number.",
    )
    doc.code_with_explain(
        "@router.delete(\"/delete/{path:path}\")\n"
        "async def delete_file(path: str):\n"
        "    file_path = _safe_path(path)\n"
        "    if not file_path.exists():\n"
        "        raise HTTPException(status_code=404, detail=\"File not found\")\n"
        "    if file_path.is_dir():\n"
        "        shutil.rmtree(file_path)\n"
        "    else:\n"
        "        file_path.unlink()\n"
        "    return {\"deleted\": str(file_path)}",
        "Same pattern: validate, then delete. If it's a directory, recursive remove; "
        "otherwise unlink (POSIX for \"remove this file\").",
    )

    doc.h3("save_api_key")
    doc.code_with_explain(
        "@router.post(\"/apikey/{project_id}\", response_model=ApiKeyResponse)\n"
        "async def save_api_key(project_id: str, api_key: str = Form(...)):\n"
        "    \"\"\"Save an API key to a file in the project folder.\n"
        "\n"
        "    ReviewPyPerAPI expects a file path to the API key, not the key itself.\n"
        "    This endpoint bridges that gap.\n"
        "    \"\"\"\n"
        "    project_dir = _safe_path(project_id)\n"
        "    if not project_dir.is_dir():\n"
        "        raise HTTPException(status_code=404, detail=\"Project not found\")\n"
        "    key_path = project_dir / \"api_key.txt\"\n"
        "    key_path.write_text(api_key.strip())\n"
        "    return ApiKeyResponse(api_key_path=str(key_path))",
        "Bridges a quirk of the upstream ReviewPyper toolkit: it wants a FILE PATH where "
        "the API key lives, not the key string itself. We write the key into "
        "<project>/api_key.txt and return the path. The user-supplied-key path "
        "(ADR-0005) hits this endpoint; the managed-credits path writes the same file from "
        "create_project.",
    )

    doc.h3("get_project_state and put_project_state")
    doc.code_with_explain(
        "@router.get(\"/projects/{project_id}/state\")\n"
        "async def get_project_state(project_id: str) -> dict[str, Any]:\n"
        "    \"\"\"Read the persisted JSON state for a project. Returns {} if missing.\"\"\"\n"
        "    project_dir = _safe_path(project_id)\n"
        "    if not project_dir.is_dir():\n"
        "        raise HTTPException(status_code=404, detail=\"Project not found\")\n"
        "    state_path = project_dir / \"state.json\"\n"
        "    if not state_path.exists():\n"
        "        return {}\n"
        "    try:\n"
        "        return json.loads(state_path.read_text())\n"
        "    except json.JSONDecodeError:\n"
        "        raise HTTPException(status_code=500, detail=\"Corrupt state.json\")",
        "Read state.json from disk and return it as parsed JSON. Missing file is fine "
        "(empty dict — common for newly created projects). Corrupt file is a server error "
        "the user can't fix in-band; they'd need to delete state.json manually. The return "
        "annotation dict[str, Any] reflects that the schema is open — the frontend's Project "
        "type is the closest thing we have to a contract.",
    )
    doc.code_with_explain(
        "@router.put(\"/projects/{project_id}/state\")\n"
        "async def put_project_state(project_id: str, state: dict[str, Any]) -> dict[str, Any]:\n"
        "    \"\"\"Replace the project state. Returns the saved state.\"\"\"\n"
        "    project_dir = _safe_path(project_id)\n"
        "    if not project_dir.is_dir():\n"
        "        raise HTTPException(status_code=404, detail=\"Project not found\")\n"
        "    (project_dir / \"state.json\").write_text(json.dumps(state, indent=2))\n"
        "    return state",
        "Replace-don't-merge: PUT semantics. The frontend always sends the FULL state when "
        "it patches a single field — the patchPipeline helper in useServerProjectState "
        "merges client-side and PUTs the result. indent=2 keeps the on-disk file readable "
        "for debugging.",
    )
    doc.callout(
        "Why state.json on disk and not a database",
        "The whole project is per-review folders under DATA_DIR — papers, OCR output, "
        "screening CSVs. Putting the state next to its files keeps a project completely "
        "self-contained: copy the folder, you've copied the project. Adding a database would "
        "split the source of truth across two systems.",
    )
    doc.page_break()


# ────── Part 3: Frontend ──────

def part3_intro(doc: Doc):
    doc.h1("Part 3 — Frontend  ·  reviewpyper_frontend/")
    doc.p(
        "The frontend is React 19 + TypeScript built with Vite, styled with Tailwind v4. "
        "Server state goes through TanStack Query (caching, retries, mutations); "
        "ephemeral client state uses useState; persistent client state uses localStorage "
        "(legacy, see ADR-0002). Configuration is fetched at boot from /config.json, which "
        "lets the same compiled bundle target dev / synthetic-test / real-API deployments by "
        "swapping the JSON file. We walk through each layer bottom-up — boot, config, HTTP, "
        "query hooks, state, components, pages."
    )


def part3_main_app(doc: Doc):
    doc.h2("Part 3.1 — main.tsx and App.tsx")

    doc.h3("main.tsx — boot")
    doc.code(
        """import { StrictMode } from 'react'
import { createRoot } from 'react-dom/client'
import './index.css'
import App from './App.tsx'
import { loadConfig } from './hooks/useConfig'
import type { AppConfig } from './hooks/useConfig'

loadConfig().then((config: AppConfig) => {
  createRoot(document.getElementById('root')!).render(
    <StrictMode>
      <App config={config} />
    </StrictMode>,
  )
}).catch((err) => {
  document.getElementById('root')!.innerHTML =
    `<div style=\"padding:2rem;color:red;font-family:monospace\">
      <h2>Failed to load config.json</h2>
      <pre>${'${err.message}'}</pre>
    </div>`
})
"""
    )
    doc.p(
        "The entry point. Three notable choices: (1) we DON'T render the app immediately — "
        "we await loadConfig() (a fetch for /config.json) first, because most of the app "
        "depends on the endpoint table. (2) We use createRoot, the React 18+ API, instead "
        "of the older ReactDOM.render. (3) We pass the config as a prop into <App>, which "
        "then puts it in a Context so any descendant can pull it via useConfig()."
    )
    doc.p(
        "<StrictMode> is a development-only React component that double-invokes effects and "
        "renders to surface unsafe patterns (use-after-unmount, derived-state effects). "
        "It's a no-op in production builds."
    )
    doc.p(
        "The catch handler renders a plain HTML error block bypassing React entirely — if "
        "config loading failed, React itself might not be initializable, so we fall back to "
        "innerHTML. document.getElementById('root')! uses TS's non-null assertion "
        "(the ! operator) — the index.html guarantees a #root element exists, so we tell "
        "the type system to trust us."
    )

    doc.h3("App.tsx — provider tree")
    doc.code(
        """import { useMemo, useState } from 'react';
import { BrowserRouter, Routes, Route } from 'react-router-dom';
import { QueryClient, QueryClientProvider, QueryCache, MutationCache } from '@tanstack/react-query';
import { ConfigProvider } from './hooks/useConfig';
import type { AppConfig } from './hooks/useConfig';
import AppLayout from './components/layout/AppLayout';
import ErrorBoundary from './components/ErrorBoundary';
import SkipLink from './components/ui/SkipLink';
import { ToastProvider, useToast } from './components/ui/Toast';
import { getErrorMessage } from './services/errorMessage';
import SetupReview from './pages/SetupReview';
import TitleScreening from './pages/TitleScreening';
import AbstractScreening from './pages/AbstractScreening';
import PdfProcessing from './pages/PdfProcessing';
import TextSections from './pages/TextSections';
import InclusionEvaluation from './pages/InclusionEvaluation';
import DataExtraction from './pages/DataExtraction';

function AppRoutes() {
  const [projectId] = useState(
    () => localStorage.getItem('reviewpyper_project_id') ?? undefined,
  );
  return (
    <Routes>
      <Route element={<AppLayout projectId={projectId} />}>
        <Route path=\"/\" element={<SetupReview />} />
        <Route path=\"/setup\" element={<SetupReview />} />
        <Route path=\"/screening\" element={<TitleScreening />} />
        <Route path=\"/abstract-screening\" element={<AbstractScreening />} />
        <Route path=\"/pdfs\" element={<PdfProcessing />} />
        <Route path=\"/text-sections\" element={<TextSections />} />
        <Route path=\"/inclusion\" element={<InclusionEvaluation />} />
        <Route path=\"/extraction\" element={<DataExtraction />} />
      </Route>
    </Routes>
  );
}
"""
    )
    doc.p(
        "AppRoutes defines the URL → component mapping. The outer <Route element={<AppLayout/>}> "
        "is a layout route — every nested route renders inside <AppLayout>'s <Outlet/>. So "
        "the sidebar, header, and pipeline stepper appear on every page automatically."
    )
    doc.p(
        "useState(() => localStorage.getItem(...)) is the LAZY initializer pattern. The "
        "function inside useState only runs on the first render; on subsequent renders, "
        "React already has the value. This pattern lets us read external mutable state "
        "(localStorage) without violating render purity — without the lazy form, every "
        "render would hit localStorage, which is not allowed during render in React 19."
    )

    doc.code(
        """function QueryProviderWithToasts({ children }: { children: React.ReactNode }) {
  const toast = useToast();
  const client = useMemo(
    () =>
      new QueryClient({
        defaultOptions: {
          queries: { retry: 1, staleTime: 30_000, refetchOnWindowFocus: false },
        },
        queryCache: new QueryCache({
          onError: (err) => toast.error(getErrorMessage(err), 'Request failed'),
        }),
        mutationCache: new MutationCache({
          onError: (err) => toast.error(getErrorMessage(err), 'Action failed'),
        }),
      }),
    [toast],
  );
  return <QueryClientProvider client={client}>{children}</QueryClientProvider>;
}
"""
    )
    doc.p(
        "Builds the TanStack Query client and wires its global error handlers to the toast "
        "system. It must live INSIDE the ToastProvider so useToast() works — that's why we "
        "have a sub-component instead of building the client at module top-level."
    )
    doc.p(
        "useMemo with [toast] as the dep array creates the client once and reuses it. "
        "Without useMemo, every render of QueryProviderWithToasts would build a fresh "
        "QueryClient, blowing the cache. retry: 1 means \"try once more on failure.\" "
        "staleTime: 30_000 (30 seconds) tells Query that fetched data is fresh for 30s — "
        "components mounting in that window get the cached data instead of re-fetching. "
        "refetchOnWindowFocus: false disables a default that's annoying for this kind of "
        "tool (the user shouldn't see a re-fetch every time they alt-tab back)."
    )

    doc.code(
        """interface AppProps { config: AppConfig; }

export default function App({ config }: AppProps) {
  return (
    <ErrorBoundary>
      <ConfigProvider value={config}>
        <ToastProvider>
          <QueryProviderWithToasts>
            <BrowserRouter>
              <SkipLink />
              <AppRoutes />
            </BrowserRouter>
          </QueryProviderWithToasts>
        </ToastProvider>
      </ConfigProvider>
    </ErrorBoundary>
  );
}
"""
    )
    doc.p(
        "The provider tree, ordered outermost-first by importance:"
    )
    doc.bullets(
        [
            "ErrorBoundary — catches uncaught render errors and shows a friendly message; "
            "everything below it is wrapped.",
            "ConfigProvider — exposes config to useConfig() everywhere.",
            "ToastProvider — exposes the toast push function to descendants.",
            "QueryProviderWithToasts — installs the TanStack Query cache.",
            "BrowserRouter — owns the URL; <Routes> uses pushState navigation under it.",
            "SkipLink — first focusable element on the page; visible only when focused; "
            "lets keyboard users jump past the sidebar nav.",
            "<AppRoutes/> — the actual page tree.",
        ]
    )
    doc.page_break()


def part3_useconfig(doc: Doc):
    doc.h2("Part 3.2 — hooks/useConfig.ts")
    doc.p(
        "Owns the runtime configuration loaded from /config.json at boot. The configuration "
        "tells the rest of the app: am I in test mode? what's the default model? where is "
        "every endpoint? Decoupling endpoints from code lets the same compiled bundle "
        "target different environments by swapping config.json."
    )
    doc.code(
        """import { createContext, useContext } from 'react';

export interface EndpointConfig { url: string; method: 'GET' | 'POST' | 'PUT' | 'DELETE'; description: string; }
export interface AppSettings { defaultModel: string; defaultReviewType: string; defaultArticleType: string; defaultPageThreshold: number; defaultSummaryType: string; }
export interface AppConfig { testMode: boolean; settings: AppSettings; endpoints: Record<string, EndpointConfig>; }

const ConfigContext = createContext<AppConfig | null>(null);
export const ConfigProvider = ConfigContext.Provider;

export function useConfig(): AppConfig {
  const config = useContext(ConfigContext);
  if (!config) throw new Error('useConfig must be used within a ConfigProvider');
  return config;
}

export function resolveEndpoint(config: AppConfig, name: string, params?: Record<string, string>): string {
  const endpoint = config.endpoints[name];
  if (!endpoint) throw new Error(`Unknown endpoint: \"${'${name}'}\"`);
  let url = endpoint.url;
  if (params) {
    for (const [key, value] of Object.entries(params)) {
      url = url.replace(`{${'${key}'}}`, encodeURIComponent(value));
    }
  }
  const isFilesOrHealth = url.includes('/files/') || url.endsWith('/health');
  if (config.testMode && !isFilesOrHealth) {
    url += (url.includes('?') ? '&' : '?') + 'test=true';
  }
  return url;
}

export async function loadConfig(): Promise<AppConfig> {
  const response = await fetch('/config.json');
  if (!response.ok) throw new Error(`Failed to load config.json: ${'${response.status}'}`);
  return response.json();
}
"""
    )
    doc.p(
        "Three TS interfaces describe the shape of config.json: per-endpoint info "
        "(url + method + description), app-wide settings (defaults), and the top-level "
        "AppConfig that holds them. Record<string, EndpointConfig> means \"object whose keys "
        "are strings and values are EndpointConfig\" — the keys are names like "
        "\"uploadFile\", \"screenTitles\"."
    )
    doc.p(
        "createContext<AppConfig | null>(null) makes a Context with default null; we throw "
        "if a consumer calls useConfig() outside the provider. ConfigProvider is just an "
        "alias for the underlying Provider component, so consumers import a friendlier name."
    )
    doc.p(
        "resolveEndpoint is the URL-builder. Given an endpoint name (\"uploadFile\") and "
        "params ({projectId: \"abc123\"}), it pulls the template URL out of the config "
        "table, substitutes {projectId} → abc123 (URL-encoded), and — if the app is in test "
        "mode and this isn't a file-router/health endpoint — appends ?test=true so the "
        "TestModeMiddleware kicks in. The /files/ exclusion is important: file uploads need "
        "to be REAL even in test mode, because the synthetic CSV the middleware writes "
        "later needs to land on a real disk path."
    )
    doc.p(
        "loadConfig is what main.tsx awaits. fetch('/config.json') hits the same origin as "
        "the app — in dev that's the Vite dev server; in prod that's nginx serving "
        "/usr/share/nginx/html/config.json. Different deployments mount different "
        "config.json files at that path."
    )
    doc.page_break()


def part3_api(doc: Doc):
    doc.h2("Part 3.3 — services/api.ts")
    doc.p(
        "The thin HTTP layer. Two namespaces (filesApi and pipelineApi) each holding a few "
        "functions that take (config, ...args) and return Promise<T>. They build a URL via "
        "resolveEndpoint, send the request via axios, and unwrap the response. No caching, "
        "no retries — that's TanStack Query's job (Part 3.4)."
    )

    doc.h3("Setup")
    doc.code_with_explain(
        "import axios from 'axios';\n"
        "import type { AppConfig } from '../hooks/useConfig';\n"
        "import { resolveEndpoint } from '../hooks/useConfig';\n"
        "import type { ProjectInfo, FileInfo, UploadResponse, ... } from '../types';\n\n"
        "const http = axios.create();",
        "One shared axios instance for the whole app. axios.create() lets us add defaults "
        "(timeout, baseURL, interceptors) here in one place if we ever need to. Today we "
        "don't — every request is fully specified at the call site.",
    )

    doc.h3("filesApi.createProject and listProjects")
    doc.code_with_explain(
        "createProject: (config: AppConfig, name: string) => {\n"
        "  const form = new FormData();\n"
        "  form.append('name', name);\n"
        "  return http.post<ProjectInfo>(resolveEndpoint(config, 'createProject'), form).then(r => r.data);\n"
        "}",
        "Build a FormData (key/value form fields), POST it. The <ProjectInfo> generic tells "
        "axios what TYPE the response data has — TS uses this to type-check downstream "
        "consumers. .then(r => r.data) drops the axios envelope so the caller just gets the "
        "JSON payload.",
    )
    doc.code_with_explain(
        "listProjects: (config: AppConfig) => http.get<ProjectInfo[]>(resolveEndpoint(config, 'listProjects')).then(r => r.data)",
        "Single-line GET. ProjectInfo[] (an array) tells the caller to expect a list.",
    )

    doc.h3("filesApi.upload — the rebuilt one")
    doc.code_with_explain(
        "upload: (\n"
        "  config: AppConfig,\n"
        "  projectId: string,\n"
        "  file: File,\n"
        "  subfolder = '',\n"
        "  onProgress?: (fraction: number) => void,\n"
        ") => {\n"
        "  const form = new FormData();\n"
        "  form.append('file', file);\n"
        "  form.append('subfolder', subfolder);\n"
        "  return http\n"
        "    .post<UploadResponse>(resolveEndpoint(config, 'uploadFile', { projectId }), form, {\n"
        "      onUploadProgress: (e) => {\n"
        "        if (onProgress && e.total) onProgress(e.loaded / e.total);\n"
        "      },\n"
        "    })\n"
        "    .then((r) => r.data);\n"
        "},",
        "The original version of this function set headers: { 'Content-Type': "
        "'multipart/form-data' } explicitly — and that was the bug. axios will set "
        "Content-Type with a unique boundary parameter automatically WHEN it sees a "
        "FormData payload, but only if the caller doesn't already have a Content-Type. By "
        "supplying \"multipart/form-data\" without a boundary, we were preventing axios from "
        "generating one, so the body became unparseable. Removing the headers option lets "
        "axios do its job. The new onProgress parameter wraps onUploadProgress (axios's "
        "callback that fires periodically with {loaded, total}). e.total can be undefined "
        "(e.g. when content-length is unknown), so we guard before dividing.",
    )
    doc.callout(
        "If a reviewer asks \"why a fraction not bytes\"",
        "Because the UI cares about percentage. A fraction (0..1) is unitless and lets the "
        "ProgressBar component multiply by 100 once. Bytes would force the consumer to know "
        "the total too.",
    )

    doc.h3("Other filesApi methods (briefly)")
    doc.bullets(
        [
            "download — GET returning blob (responseType: 'blob'); the blob can be given to a "
            "<a download> anchor or saved via a library.",
            "list — GET; passes subfolder as a query param via { params }.",
            "delete — DELETE; URL has a {path} that can contain slashes.",
            "saveApiKey — POST a FormData with api_key field.",
            "getProjectState / putProjectState — used by useServerProjectState (the planned "
            "future state hook).",
        ]
    )

    doc.h3("pipelineApi")
    doc.code_with_explain(
        "export const pipelineApi = {\n"
        "  screenTitles: (config: AppConfig, params: Record<string, unknown>) => http.post<TitleScreenResponse>(resolveEndpoint(config, 'screenTitles'), params).then(r => r.data),\n"
        "  screenAbstracts: (config: AppConfig, params: Record<string, unknown>) => http.post<AbstractScreenResponse>(resolveEndpoint(config, 'screenAbstracts'), params).then(r => r.data),\n"
        "  // ...\n"
        "};",
        "One function per pipeline endpoint. Each takes a params object (validated upstream "
        "by Zod schemas in useApi.ts), POSTs it as JSON, returns the typed response. "
        "Record<string, unknown> here is intentionally loose because the validated shape is "
        "enforced one layer up.",
    )
    doc.page_break()


def part3_useapi(doc: Doc):
    doc.h2("Part 3.4 — hooks/useApi.ts")
    doc.p(
        "TanStack Query wrappers around services/api.ts. Each pipeline action becomes a "
        "useMutation hook; lists become useQuery hooks. The wrappers also pull config out "
        "of context (so callers don't have to pass it) and run params through Zod schemas "
        "before sending."
    )

    doc.h3("List queries — useServerProjects, useProjectFiles")
    doc.code_with_explain(
        "export function useServerProjects() {\n"
        "  const config = useConfig();\n"
        "  return useQuery({ queryKey: ['server-projects'], queryFn: () => filesApi.listProjects(config) });\n"
        "}",
        "useQuery with key ['server-projects']. Any consumer of this hook gets the same "
        "cached data (deduplicated) and any mutation that calls "
        "queryClient.invalidateQueries({queryKey:['server-projects']}) triggers a re-fetch.",
    )

    doc.h3("Mutation hooks — useCreateServerProject pattern")
    doc.code_with_explain(
        "export function useCreateServerProject() {\n"
        "  const config = useConfig();\n"
        "  const qc = useQueryClient();\n"
        "  return useMutation({\n"
        "    mutationFn: (name: string) => filesApi.createProject(config, name),\n"
        "    onSuccess: () => qc.invalidateQueries({ queryKey: ['server-projects'] }),\n"
        "  });\n"
        "}",
        "Turns the POST into a UI-friendly mutation. The hook returns "
        "{ mutate, mutateAsync, isPending, isError, error, data, ... }. mutateAsync "
        "returns a Promise so callers can await it; mutate is fire-and-forget. "
        "On success we invalidate the projects list so a project picker stays accurate.",
    )

    doc.h3("useUploadFile — wired to onProgress")
    doc.code_with_explain(
        "export function useUploadFile() {\n"
        "  const config = useConfig();\n"
        "  return useMutation({\n"
        "    mutationFn: ({ projectId, file, subfolder, onProgress }: { ... }) =>\n"
        "      filesApi.upload(config, projectId, file, subfolder, onProgress),\n"
        "  });\n"
        "}",
        "Plumbing the progress callback all the way through: the page passes onProgress in "
        "the mutate call, this hook destructures it from the variables and passes it on to "
        "filesApi.upload, which passes it to axios. Caller controls progress reporting from "
        "one place.",
    )

    doc.h3("Pipeline mutation hooks (representative)")
    doc.code_with_explain(
        "export function useScreenTitles() {\n"
        "  const config = useConfig();\n"
        "  return useMutation({\n"
        "    mutationFn: (params: Record<string, unknown>) =>\n"
        "      pipelineApi.screenTitles(config, validateOrThrow(ScreenTitlesParams, params)),\n"
        "  });\n"
        "}",
        "validateOrThrow runs the params through a Zod schema (defined in services/schemas.ts) "
        "and either returns the validated, typed object or throws if a required field is "
        "missing or wrong type. This catches frontend bugs before they hit the network. The "
        "remaining nine pipeline hooks (useScreenAbstracts, useDownloadPdfs, ...) are all "
        "this exact same shape with different schemas — a candidate for deduplication into "
        "a single factory in a follow-up refactor.",
    )
    doc.page_break()


def part3_useprojectstate(doc: Doc):
    doc.h2("Part 3.5 — hooks/useProjectState.ts")
    doc.p(
        "The legacy localStorage-backed project store. Per ADR-0002, this is deprecated in "
        "favor of useServerProjectState — but we haven't migrated the pages yet. It still "
        "owns the source of truth in the running app today."
    )

    doc.h3("Module load")
    doc.code_with_explain(
        "import { useState, useCallback } from 'react';\n"
        "import type { Project, PipelineState, LLMConfig } from '../types';\n\n"
        "const STORAGE_KEY = 'reviewpyper_projects';",
        "Imports React's useState and useCallback (memoize a function reference across "
        "renders); types are pulled from the central types/index.ts. The storage key is the "
        "single localStorage entry — its value is a JSON-encoded object: { [id]: Project }.",
    )

    doc.h3("Persistence helpers")
    doc.code_with_explain(
        "function loadProjects(): Record<string, Project> {\n"
        "  try {\n"
        "    const raw = localStorage.getItem(STORAGE_KEY);\n"
        "    return raw ? JSON.parse(raw) : {};\n"
        "  } catch { return {}; }\n"
        "}\n"
        "function saveProjects(projects: Record<string, Project>) {\n"
        "  localStorage.setItem(STORAGE_KEY, JSON.stringify(projects));\n"
        "}",
        "loadProjects swallows JSON-parse errors and corrupt storage by returning {}. "
        "Empirically, you do see corrupt entries occasionally (someone hand-edits storage, "
        "browser extension messes with it) and the application should keep working.",
    )

    doc.h3("The hook")
    doc.code_with_explain(
        "export function useProjectState() {\n"
        "  const [projects, setProjects] = useState<Record<string, Project>>(loadProjects);",
        "useState's initialState argument is a FUNCTION here (loadProjects), not a value. "
        "React calls it once on mount. This is the lazy initializer pattern — cheap to set "
        "up, expensive function only runs once.",
    )
    doc.code_with_explain(
        "  const getProject = useCallback((id: string): Project | undefined => projects[id], [projects]);",
        "Returns the Project for a given id or undefined. useCallback ensures the returned "
        "function reference stays stable as long as projects doesn't change — useful so "
        "downstream useEffect dep arrays don't re-run unnecessarily.",
    )
    doc.code_with_explain(
        "  const createProject = useCallback((id, name, researchQuestion, reviewType) => {\n"
        "    const project: Project = { id, name, research_question: researchQuestion, review_type: reviewType, created_at: new Date().toISOString(), pipeline_state: {} };\n"
        "    setProjects(prev => {\n"
        "      const updated = { ...prev, [id]: project };\n"
        "      saveProjects(updated);\n"
        "      return updated;\n"
        "    });\n"
        "    return project;\n"
        "  }, []);",
        "Creates a Project object, stores it in both state AND localStorage. Note the "
        "functional setProjects(prev => ...) form — we read prev rather than the closed-over "
        "projects, which avoids a stale-closure bug if multiple updates fire in the same "
        "frame. saveProjects runs INSIDE the setter callback so on-disk and in-state stay "
        "consistent. We persist before returning.",
    )
    doc.code_with_explain(
        "  const updatePipelineState = useCallback((id: string, patch: Partial<PipelineState>) => {\n"
        "    setProjects(prev => {\n"
        "      const project = prev[id];\n"
        "      if (!project) return prev;\n"
        "      const updated = { ...prev, [id]: { ...project, pipeline_state: { ...project.pipeline_state, ...patch } } };\n"
        "      saveProjects(updated);\n"
        "      return updated;\n"
        "    });\n"
        "  }, []);",
        "Patch-merge a Partial<PipelineState> into an existing project's pipeline_state. "
        "Three nested spreads: outer copies the projects map, mid copies the project, "
        "inner copies pipeline_state with patch overrides applied. If the project doesn't "
        "exist (prev[id] is undefined), return prev unchanged — same reference, so React "
        "skips a re-render.",
    )
    doc.callout(
        "Why this entire hook is doomed",
        "ADR-0002 — the source of truth needs to be the server's state.json. Two browsers "
        "looking at the same project see different states because they each read from their "
        "own localStorage. The migration is intentionally incremental.",
    )
    doc.page_break()


def part3_fileupload(doc: Doc):
    doc.h2("Part 3.6 — components/ui/FileUpload.tsx")
    doc.p(
        "The drag-and-drop dropzone. The file we rewrote. The original version had three "
        "issues: misleading drag-and-drop messaging without actual handlers, an unused "
        "maxSizeMB prop, and an opaque \"Uploading...\" spinner with no progress."
    )

    doc.h3("Type signature")
    doc.code_with_explain(
        "interface FileUploadProps {\n"
        "  accept?: string;\n"
        "  multiple?: boolean;\n"
        "  onFiles: (files: File[]) => void;\n"
        "  label?: string;\n"
        "  hint?: string;\n"
        "  maxSizeMB?: number;\n"
        "  disabled?: boolean;\n"
        "}",
        "All optional except onFiles. accept maps to the HTML input's accept attribute "
        "(e.g. \".csv\"). multiple lets the user pick more than one file. onFiles receives "
        "the validated File[] array. maxSizeMB triggers client-side size validation. "
        "disabled greys out the dropzone (e.g. while an upload is in flight).",
    )

    doc.h3("formatBytes helper")
    doc.code_with_explain(
        "function formatBytes(bytes: number): string {\n"
        "  if (bytes < 1024) return `${'${bytes}'} B`;\n"
        "  if (bytes < 1024 ** 2) return `${'${(bytes / 1024).toFixed(0)}'} KB`;\n"
        "  if (bytes < 1024 ** 3) return `${'${(bytes / 1024 ** 2).toFixed(1)}'} MB`;\n"
        "  return `${'${(bytes / 1024 ** 3).toFixed(2)}'} GB`;\n"
        "}",
        "Human-friendly size formatter for the error message when a file is too big. "
        "Could be one line via Intl.NumberFormat with units, but this is clearer at a "
        "glance.",
    )

    doc.h3("Component setup")
    doc.code_with_explain(
        "export default function FileUpload({\n"
        "  accept, multiple, onFiles,\n"
        "  label = 'Drop files here or click to browse',\n"
        "  hint, maxSizeMB, disabled = false,\n"
        "}: FileUploadProps) {\n"
        "  const inputRef = useRef<HTMLInputElement>(null);\n"
        "  const [dragActive, setDragActive] = useState(false);\n"
        "  const [error, setError] = useState<string | null>(null);",
        "Default values for label/disabled in the destructuring. inputRef gives us a handle "
        "to the hidden <input type=\"file\"> so we can programmatically open the OS file "
        "picker (calling .click() on it). dragActive drives the visual highlight when "
        "someone drags over the zone. error holds a user-facing validation message.",
    )

    doc.h3("validate — the real maxSizeMB use")
    doc.code_with_explain(
        "  const validate = useCallback(\n"
        "    (files: File[]): { ok: File[]; reason?: string } => {\n"
        "      if (!files.length) return { ok: [] };\n"
        "      if (maxSizeMB) {\n"
        "        const limit = maxSizeMB * 1024 * 1024;\n"
        "        const oversized = files.find((f) => f.size > limit);\n"
        "        if (oversized) {\n"
        "          return {\n"
        "            ok: [],\n"
        "            reason: `${'${oversized.name}'} is ${'${formatBytes(oversized.size)}'} — exceeds the ${'${maxSizeMB}'} MB limit.`,\n"
        "          };\n"
        "        }\n"
        "      }\n"
        "      return { ok: files };\n"
        "    },\n"
        "    [maxSizeMB],\n"
        "  );",
        "Returns either { ok: validFiles } or { ok: [], reason: 'message' }. The struct-y "
        "return type makes the caller pattern-match on success vs. failure without a "
        "separate boolean. Find the first oversized file by size in bytes, complain with a "
        "human-readable message that includes the actual size and the limit.",
    )

    doc.h3("handleFiles — pipe through validation")
    doc.code_with_explain(
        "  const handleFiles = useCallback(\n"
        "    (files: File[]) => {\n"
        "      setError(null);\n"
        "      const result = validate(files);\n"
        "      if (result.reason) {\n"
        "        setError(result.reason);\n"
        "        return;\n"
        "      }\n"
        "      if (result.ok.length) onFiles(result.ok);\n"
        "    },\n"
        "    [validate, onFiles],\n"
        "  );",
        "Single funnel: every entry point (file picker, drop) calls this. Clear any prior "
        "error first so a successful pick clears the validation message. If validation "
        "failed, store the message; otherwise hand the files to the parent.",
    )

    doc.h3("Drag handlers — the missing piece")
    doc.code_with_explain(
        "  const handleDragOver = useCallback((e: React.DragEvent) => {\n"
        "    e.preventDefault();\n"
        "    if (!disabled) setDragActive(true);\n"
        "  }, [disabled]);\n\n"
        "  const handleDragLeave = useCallback((e: React.DragEvent) => {\n"
        "    e.preventDefault();\n"
        "    setDragActive(false);\n"
        "  }, []);\n\n"
        "  const handleDrop = useCallback((e: React.DragEvent) => {\n"
        "    e.preventDefault();\n"
        "    setDragActive(false);\n"
        "    if (disabled) return;\n"
        "    const files = Array.from(e.dataTransfer.files);\n"
        "    handleFiles(files);\n"
        "  }, [disabled, handleFiles]);",
        "preventDefault on dragover is REQUIRED — the default browser behavior when you "
        "drop a file on a webpage is to navigate to it (open the file). preventDefault "
        "tells the browser not to. We toggle dragActive so the dropzone can highlight, "
        "and on drop we read e.dataTransfer.files (a FileList, similar to NodeList — "
        "iterable but not an Array, hence Array.from).",
    )

    doc.h3("Click and keyboard activation")
    doc.code_with_explain(
        "  const handleClick = useCallback(() => {\n"
        "    if (!disabled) inputRef.current?.click();\n"
        "  }, [disabled]);\n\n"
        "  const handleKeyDown = useCallback((e: React.KeyboardEvent) => {\n"
        "    if ((e.key === 'Enter' || e.key === ' ') && !disabled) {\n"
        "      e.preventDefault();\n"
        "      inputRef.current?.click();\n"
        "    }\n"
        "  }, [disabled]);",
        "The dropzone is a <div role=\"button\"> not an actual <button>, so we have to "
        "implement keyboard activation by hand: Enter and Space both fire the file picker "
        "(matching native button behavior). preventDefault on Space prevents the page from "
        "scrolling.",
    )

    doc.h3("JSX — the dropzone div")
    doc.code(
        """  return (
    <div className=\"space-y-2\">
      <div
        role=\"button\"
        tabIndex={disabled ? -1 : 0}
        aria-disabled={disabled}
        aria-describedby={hint || error ? 'file-upload-hint' : undefined}
        onClick={handleClick}
        onKeyDown={handleKeyDown}
        onDragOver={handleDragOver}
        onDragEnter={handleDragOver}
        onDragLeave={handleDragLeave}
        onDrop={handleDrop}
        className={cn(
          'flex flex-col items-center justify-center w-full h-36 border-2 border-dashed rounded-lg transition-colors outline-none',
          'focus-visible:ring-2 focus-visible:ring-primary-500 focus-visible:ring-offset-2',
          disabled && 'opacity-50 cursor-not-allowed border-gray-300',
          !disabled && dragActive && 'border-primary-500 bg-primary-50',
          !disabled && !dragActive && 'border-gray-300 hover:border-primary-400 hover:bg-primary-50/40 cursor-pointer',
          error && 'border-danger-400',
        )}
      >
"""
    )
    doc.p(
        "role=\"button\" announces this region as a button to screen readers. tabIndex={0} "
        "makes it focusable; tabIndex={-1} when disabled removes it from the keyboard tab "
        "order. aria-describedby links the dropzone to the hint or error message below. "
        "The className collapses through cn() — a tailwind-merge helper that resolves "
        "conflicts (hover:border-primary-400 vs border-primary-500) by latest-wins. The "
        "five conditional class blocks express the visual state machine: disabled, dragging "
        "over, idle, error."
    )

    doc.code(
        """        <Upload className={cn('h-8 w-8 mb-2', dragActive ? 'text-primary-600' : 'text-gray-400')} aria-hidden=\"true\" />
        <span className=\"text-sm font-medium text-gray-700\">{label}</span>
        {hint && (
          <span id=\"file-upload-hint\" className=\"text-xs text-gray-500 mt-1\">
            {hint}
            {maxSizeMB ? ` · max ${'${maxSizeMB}'} MB` : ''}
          </span>
        )}
        <input
          ref={inputRef}
          type=\"file\"
          className=\"sr-only\"
          accept={accept}
          multiple={multiple}
          onChange={handleChange}
          disabled={disabled}
          tabIndex={-1}
        />
      </div>
      {error && (
        <p role=\"alert\" className=\"flex items-center gap-1.5 text-xs text-danger-700\">
          <FileWarning className=\"h-3.5 w-3.5 flex-shrink-0\" aria-hidden=\"true\" />
          {error}
        </p>
      )}
    </div>
  );
}
"""
    )
    doc.p(
        "The Upload icon recolors when dragActive. The hint span shows the file-type hint "
        "plus, if maxSizeMB is set, a \"max 500 MB\" suffix — telling the user the limit "
        "BEFORE they try to upload. The hidden <input> is the actual file picker; "
        "className=\"sr-only\" makes it visually hidden but still functional. The error <p> "
        "uses role=\"alert\" so screen readers announce it immediately when it appears."
    )
    doc.page_break()


def part3_stepper(doc: Doc):
    doc.h2("Part 3.7 — components/PipelineStepper.tsx")
    doc.p(
        "Horizontal step indicator showing all seven pipeline stages, with completion "
        "derived from the project's PipelineState. Lives in the header of every page."
    )

    doc.h3("Step table")
    doc.code(
        """const STEPS: Step[] = [
  { id: 1, label: 'Setup',     short: 'Setup',      path: '/setup',              isComplete: (p) => !!p.api_key_path },
  { id: 2, label: 'Titles',    short: 'Titles',     path: '/screening',          isComplete: (p) => !!p.title_output_csv_path },
  { id: 3, label: 'Abstracts', short: 'Abstracts',  path: '/abstract-screening', isComplete: (p) => !!p.master_list_path },
  { id: 4, label: 'PDFs',      short: 'PDFs',       path: '/pdfs',               isComplete: (p) => !!p.ocr_output_dir },
  { id: 5, label: 'Sections',  short: 'Sections',   path: '/text-sections',      isComplete: (p) => !!p.json_dir },
  { id: 6, label: 'Inclusion', short: 'Inclusion',  path: '/inclusion',          isComplete: (p) => !!p.inclusion_automated_csv_path },
  { id: 7, label: 'Extraction',short: 'Extract',    path: '/extraction',         isComplete: (p) => !!p.extraction_automated_csv_path },
];
"""
    )
    doc.p(
        "Static array; building it once at module top-level. Each step knows its own "
        "completion predicate (isComplete) — we just check whether the corresponding "
        "PipelineState path exists. !!x converts to a boolean (truthy non-empty path → "
        "true, undefined / empty → false)."
    )

    doc.h3("Component")
    doc.code(
        """export default function PipelineStepper({ pipeline }: { pipeline: PipelineState }) {
  const navigate = useNavigate();
  const location = useLocation();

  const activeIndex = STEPS.findIndex(
    (s) => location.pathname === s.path || (s.path === '/setup' && location.pathname === '/'),
  );
  ...
"""
    )
    doc.p(
        "Reads current URL via useLocation (from react-router-dom). activeIndex is the "
        "index of the step whose path matches — also handling the special case where "
        "location.pathname is \"/\" (which routes to SetupReview)."
    )
    doc.p(
        "Each step button is wrapped in <li>. The styling is a 3-state machine per step: "
        "done (filled circle with checkmark + connector painted), current (highlighted but "
        "no checkmark), or future (greyed). reachable (computed in render) gates whether "
        "the user can click forward — you can't jump from Setup to Inclusion if you "
        "haven't completed the in-between steps. The connector line between steps adopts "
        "the primary color when the previous step is done."
    )
    doc.page_break()


def part3_setupreview(doc: Doc):
    doc.h2("Part 3.8 — pages/SetupReview.tsx")
    doc.p(
        "Step 1 of the pipeline: define the review and configure API access. The page was "
        "rewritten to remove the derived-state useEffect anti-pattern. An earlier draft "
        "exposed an OpenAI / Anthropic radio switcher; that was removed because the "
        "backend currently only auto-saves the OpenAI key — the provider switch was "
        "advertising a capability that wasn't end-to-end supported. The page now defaults "
        "to OpenAI silently."
    )

    doc.h3("Setup phase — lazy init from existing")
    doc.code_with_explain(
        "  const [projectId, setProjectId] = useState<string | null>(() =>\n"
        "    localStorage.getItem(PROJECT_ID_STORAGE_KEY),\n"
        "  );\n"
        "  const existing = projectId ? getProject(projectId) : undefined;\n\n"
        "  const [name, setName] = useState(existing?.name ?? '');\n"
        "  const [question, setQuestion] = useState(existing?.research_question ?? '');\n"
        "  const [reviewType, setReviewType] = useState(existing?.review_type ?? settings.defaultReviewType);\n"
        "  const [userApiKey, setUserApiKey] = useState('');",
        "All initial values come from the EXISTING project (if any). useState's lazy-init "
        "function form means localStorage is read once at mount, not on every render. The "
        "form fields are then independently editable. This is the corrected pattern; the "
        "previous version had a useEffect that copied existing into state on mount, which "
        "is the textbook \"derived state in effect\" anti-pattern.",
    )

    doc.h3("handleCreateProject")
    doc.code_with_explain(
        "  const handleCreateProject = async () => {\n"
        "    const serverProject = await createServerProject.mutateAsync(name);\n"
        "    const id = serverProject.project_id;\n"
        "    createProject(id, name, question, reviewType as 'standard' | 'rapid' | 'scoping');\n"
        "    setProjectId(id);\n"
        "    localStorage.setItem(PROJECT_ID_STORAGE_KEY, id);",
        "Sequence: hit the API to create a server-side project (returns the ID + key path); "
        "mirror it in localStorage via createProject (the legacy hook); store the active ID "
        "for navigation. The two stores temporarily diverge — see ADR-0002 — but for the "
        "lifetime of this session they're consistent.",
    )
    doc.code_with_explain(
        "    if (serverProject.api_key_path) {\n"
        "      setLLMConfig(id, {\n"
        "        provider: 'openai',\n"
        "        api_key: '(server-managed)',\n"
        "        model: settings.defaultModel,\n"
        "        model_choice: settings.defaultModel,\n"
        "      });\n"
        "      updatePipelineState(id, { api_key_path: serverProject.api_key_path });\n"
        "    } else if (userApiKey) {\n"
        "      const result = await saveApiKey.mutateAsync({ projectId: id, apiKey: userApiKey });\n"
        "      ...\n"
        "    }\n"
        "  };",
        "Two branches per ADR-0005: managed credits (server set OPENAI_API_KEY → api_key.txt "
        "already on disk) or user-supplied (POST the key, server saves it). Either way, "
        "we record api_key_path in PipelineState so downstream pipeline calls have it. "
        "Provider is hardcoded to 'openai' since that is the only provider the backend "
        "currently auto-saves a key for.",
    )
    doc.callout(
        "Why the provider switcher was removed",
        "The backend's create_project endpoint reads OPENAI_API_KEY from the environment "
        "and writes it to api_key.txt. There's no parallel ANTHROPIC_API_KEY pickup, so "
        "advertising both providers in the UI promised a capability the system wasn't "
        "actually delivering end-to-end. Until ReviewPyper's pipeline routers natively "
        "support Anthropic, the page sticks with one provider.",
    )
    doc.page_break()


def part3_pdf(doc: Doc):
    doc.h2("Part 3.10 — pages/PdfProcessing.tsx (OA retrieval mode)")
    doc.p(
        "Step 4 of the pipeline. The first sub-step downloads PDFs for every paper in the "
        "screened master list. ReviewPyper's BulkPDFDownloaderV2 uses pypaperretriever, "
        "which queries Unpaywall and PubMed Central for legally-free Open Access copies; "
        "if no OA copy is found and Sci-Hub fallback is allowed, it tries Sci-Hub. Recent "
        "UX work made the OA-only mode the prominent default."
    )

    doc.h3("RetrievalMode union")
    doc.code_with_explain(
        "type RetrievalMode = 'oa' | 'oa_plus_scihub';\n\n"
        "const RETRIEVAL_MODES: { value: RetrievalMode; label: string; description: string; icon: typeof Globe2 }[] = [\n"
        "  { value: 'oa', label: 'Open Access only', description: 'PubMed + Unpaywall. Free, legal everywhere.', icon: Globe2 },\n"
        "  { value: 'oa_plus_scihub', label: 'Include Sci-Hub fallback', description: 'Try OA first, then Sci-Hub for paywalled papers.', icon: Library },\n"
        "];",
        "Two-mode segmented control. The first option is the safe default — Unpaywall + "
        "PubMed Central return only legally-free copies, so users who run this from "
        "institutional networks face no compliance worry. The second option enables the "
        "Sci-Hub fallback for paywalled papers; it's offered for completeness but isn't "
        "the default. Backend-wise, both modes hit the same endpoint with allow_scihub "
        "true or false.",
    )

    doc.h3("Mode → backend boolean")
    doc.code_with_explain(
        "  const [retrievalMode, setRetrievalMode] = useState<RetrievalMode>('oa');\n"
        "  const allowScihub = retrievalMode === 'oa_plus_scihub';",
        "Derived boolean in render — no useEffect needed. allowScihub is the value "
        "actually sent to the API; the segmented control is just UX sugar over a single "
        "binary flag.",
    )
    doc.callout(
        "What's already there for free",
        "OA retrieval doesn't need new server code. ReviewPyper already calls Unpaywall "
        "and PubMed Central when allow_scihub=False; the only thing the recent UX change "
        "did was make this the obvious default. The button label even adapts to read "
        "\"Retrieve OA PDFs\" in OA mode and \"Retrieve PDFs\" otherwise.",
    )
    doc.page_break()


def part3_titlescreening(doc: Doc):
    doc.h2("Part 3.9 — pages/TitleScreening.tsx")
    doc.p(
        "Step 2 — upload references CSV, run AI screening on the titles. Same shape "
        "applies to the other six pipeline pages; this is the canonical example."
    )

    doc.h3("State setup")
    doc.code_with_explain(
        "  const [projectId] = useState(() => localStorage.getItem('reviewpyper_project_id') ?? '');\n"
        "  const { getProject, getPipelineState, updatePipelineState } = useProjectState();\n"
        "  const project = getProject(projectId);\n"
        "  const pipeline = getPipelineState(projectId);\n"
        "  const [keywords, setKeywords] = useState('');\n"
        "  const [isScreening, setIsScreening] = useState(false);\n"
        "  const [uploadedPath, setUploadedPath] = useState(pipeline.csv_path ?? '');\n"
        "  const [outputPath, setOutputPath] = useState(pipeline.title_output_csv_path ?? '');\n"
        "  const [uploadFraction, setUploadFraction] = useState<number | undefined>(undefined);\n"
        "  const [uploadedFile, setUploadedFile] = useState<{ name: string; size: number } | null>(null);\n"
        "  const uploadFile = useUploadFile();\n"
        "  const screenTitles = useScreenTitles();",
        "Lazy init from localStorage; pull current project + pipeline state via hook; "
        "ten state pieces for everything the UI tracks (keywords input, screening progress, "
        "uploaded/output paths, upload fraction for the progress bar, uploaded-file metadata "
        "for display, plus the two TanStack mutations).",
    )

    doc.h3("handleUpload — wired to onProgress")
    doc.code_with_explain(
        "  const handleUpload = useCallback(async (files: File[]) => {\n"
        "    if (!files.length || !projectId) return;\n"
        "    const file = files[0];\n"
        "    setUploadFraction(0);\n"
        "    setUploadedFile({ name: file.name, size: file.size });\n"
        "    try {\n"
        "      const result = await uploadFile.mutateAsync({\n"
        "        projectId,\n"
        "        file,\n"
        "        onProgress: (f) => setUploadFraction(f),\n"
        "      });\n"
        "      setUploadedPath(result.path);\n"
        "      updatePipelineState(projectId, { csv_path: result.path });\n"
        "    } finally {\n"
        "      setUploadFraction(undefined);\n"
        "    }\n"
        "  }, [projectId, uploadFile, updatePipelineState]);",
        "Receives File[] from the FileUpload component. Reset progress to 0; capture the "
        "filename + size for the success message; await the mutation with onProgress wired "
        "to setUploadFraction. The finally block clears uploadFraction so the ProgressBar "
        "unmounts whether the upload succeeded or failed. updatePipelineState writes the "
        "uploaded server path into PipelineState so the next step (screen) can read it.",
    )

    doc.h3("ProgressBar wiring")
    doc.code(
        """          {uploadFile.isPending && uploadedFile && (
            <div className=\"mt-4\" role=\"status\" aria-live=\"polite\">
              <ProgressBar
                value={uploadFraction}
                label={`Uploading ${'${uploadedFile.name}'} · ${'${
                  uploadFraction !== undefined ? `${'${Math.round(uploadFraction * 100)}'}%` : 'preparing…'
                }'}`}
              />
            </div>
          )}
"""
    )
    doc.p(
        "Render the progress bar only while a mutation is pending and we have file metadata. "
        "role=\"status\" + aria-live=\"polite\" tells screen readers to announce updates "
        "without interrupting (vs. \"alert\" which is interrupting). The label is a "
        "template literal that switches between \"preparing…\" (before the first progress "
        "event) and \"42%\" once we have a fraction."
    )
    doc.page_break()


def part4(doc: Doc):
    doc.h1("Part 4 — Build infrastructure")

    doc.h2("Part 4.1 — Dockerfile.prod")
    doc.code(
        """FROM python:3.11-slim
ARG DEBIAN_FRONTEND=noninteractive
WORKDIR /app

RUN apt-get update \\
 && apt-get install -y --no-install-recommends \\
    git \\
    tesseract-ocr \\
    poppler-utils \\
    libgl1 \\
 && rm -rf /var/lib/apt/lists/*

RUN git clone https://github.com/Calvinwhow/ReviewPyper.git /app/ReviewPyper \\
 && if [ ! -f /app/ReviewPyper/__init__.py ]; then touch /app/ReviewPyper/__init__.py; fi

COPY reviewpyper_api /app/reviewpyper_api

# textract dependency conflict workaround …
RUN sed '/^textract/Id' /app/ReviewPyper/requirements.txt > /tmp/reqs.txt \\
 && pip install --no-cache-dir -r /tmp/reqs.txt \\
 && pip install --no-cache-dir --no-deps textract==1.6.5 \\
 && pip install --no-cache-dir \\
    'argcomplete~=1.10.0' 'chardet==3.*' 'docx2txt~=0.8' \\
    'extract-msg<0.30' 'pdfminer.six==20191110' \\
    'python-pptx~=0.6.18' 'six~=1.12.0' 'SpeechRecognition~=3.8.1' \\
    'xlrd~=1.2.0' \\
 && pip install --no-cache-dir fastapi 'uvicorn[standard]' python-multipart

ENV PYTHONPATH=/app/reviewpyper_api:/app:/app/ReviewPyper
ENV DATA_DIR=/data
RUN mkdir -p /data
WORKDIR /app/reviewpyper_api
EXPOSE 8000
CMD [\"uvicorn\", \"main:app\", \"--host\", \"0.0.0.0\", \"--port\", \"8000\"]
"""
    )
    doc.p(
        "Walks bottom-up of cost: install OS deps, clone ReviewPyper from GitHub at build "
        "time (we don't fork — we always pull latest at image build), copy our API code, "
        "deal with the textract dependency mess (textract pins beautifulsoup4~=4.8 but "
        "pypaperretriever requires 4.12; we strip textract from the upstream requirements, "
        "install everything else first, then install textract code-only with --no-deps and "
        "manually add the runtime deps). Set PYTHONPATH so main.py's import ReviewPyper "
        "works without sys.path mucking. EXPOSE 8000 documents the listening port. "
        "uvicorn boots FastAPI."
    )
    doc.callout(
        "Why git clone instead of pip install",
        "ReviewPyper isn't on PyPI. We could fork and vendor a copy, but cloning at build "
        "time means we always pick up the upstream's latest fixes. The trade-off is "
        "non-reproducible builds; if upstream pushes a breaking change, our next image "
        "build will inherit it. Pinning the commit SHA in the clone command would fix that.",
    )

    doc.h2("Part 4.2 — nginx.conf")
    doc.code(
        """server {
    listen 80;
    root /usr/share/nginx/html;
    index index.html;

    client_max_body_size 500m;

    # SPA fallback — all routes serve index.html
    location / {
        try_files $uri $uri/ /index.html;
    }

    # Proxy API calls to the ReviewPyper API
    location /api/ {
        proxy_pass http://reviewpyper-api:8000/;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_read_timeout 600s;
        proxy_request_buffering off;
    }
}
"""
    )
    doc.p(
        "The single nginx server block does two things: serve the SPA (with the try_files "
        "fallback to index.html so client-side routes like /screening don't 404), and "
        "reverse-proxy /api/* to the API container. client_max_body_size 500m caps "
        "upload size at 500 MB — must match the frontend's MAX_UPLOAD_MB. proxy_read_timeout "
        "extends the default 60s to 600s for long pipeline calls. proxy_request_buffering "
        "off is the magic line for big uploads: by default nginx buffers the request body "
        "to disk before forwarding, which doubles disk I/O and adds latency; off makes "
        "nginx stream the body straight to FastAPI."
    )

    doc.h2("Part 4.3 — vite.config.ts")
    doc.code(
        """import { defineConfig } from 'vite'
import react from '@vitejs/plugin-react'
import tailwindcss from '@tailwindcss/vite'

export default defineConfig({
  plugins: [react(), tailwindcss()],
  server: {
    allowedHosts: true,
    proxy: {
      '/api': {
        target: process.env.VITE_API_PROXY || 'http://localhost:8000',
        rewrite: (path) => path.replace(/^\\/api/, ''),
      }
    }
  },
  test: {
    globals: true,
    environment: 'jsdom',
    setupFiles: ['./src/test/setup.ts'],
    css: false,
    include: ['src/**/*.{test,spec}.{ts,tsx}'],
  },
})
"""
    )
    doc.p(
        "Two Vite plugins: react() handles JSX/TSX transformation and Fast Refresh (hot "
        "module replacement preserving component state); tailwindcss() processes Tailwind "
        "v4 directives in CSS. The dev server proxies /api/* to localhost:8000 (or "
        "wherever VITE_API_PROXY points). The rewrite strips the /api prefix because "
        "FastAPI doesn't expect it — it expects /files/upload/..., not /api/files/upload/.... "
        "(In production, nginx does the same strip via proxy_pass http://api:8000/.) "
        "The test block configures Vitest: jsdom for DOM emulation, global expect/test "
        "imports, MSW setup file."
    )

    doc.h2("Part 4.4 — index.css")
    doc.p(
        "Tailwind v4 CSS-first design tokens. The @theme block defines the editorial "
        "palette (warm-paper background, deep-teal primary), three-font typography "
        "(Source Serif 4 / IBM Plex Sans / JetBrains Mono), and base styles. Walking "
        "through the highlights:"
    )
    doc.code(
        """@import \"tailwindcss\";

@theme {
  --color-background: oklch(98.4% 0.006 85);
  --color-foreground: oklch(18% 0.01 264);
  --color-primary: oklch(38% 0.07 200);
  ...
  --font-display: 'Source Serif 4', ui-serif, Georgia, serif;
  --font-sans: 'IBM Plex Sans', ui-sans-serif, system-ui, ...;
  --font-mono: 'JetBrains Mono', ui-monospace, ...;
}
"""
    )
    doc.p(
        "OKLCH is a perceptual color space — equal numerical changes look like equal "
        "perceptual changes (unlike HSL, where the same hue change can vary in apparent "
        "brightness). Three numbers: lightness (0-100%), chroma (0+), hue angle (0-360°). "
        "The semantic tokens (--color-foreground, --color-primary) are what components "
        "use; the legacy primary-50…950 scale is preserved for any straggler usages."
    )
    doc.p(
        "@media (prefers-reduced-motion: reduce) wraps a snippet that crushes all "
        "animation durations to ~0ms — respecting users with vestibular disorders."
    )
    doc.page_break()


def part6_validation(doc: Doc):
    doc.h1("Part 6 — Validation harness  ·  validation/")
    doc.p(
        "The validation/ directory is a standalone testing harness that drives the running "
        "ReviewPyperAPI against published Cochrane benchmarks (CLEF-TAR 2017–2019 and "
        "Cohen 2006) and reports recall, precision, F1, and WSS@95 against the gold "
        "standard. None of this is wired into the application's runtime — it's a research "
        "tool that talks to the same HTTP API a browser does."
    )

    doc.h2("Layout")
    doc.code(
        """validation/
├── README.md                    overview of datasets and topic recommendations
├── inventory.csv                246 rows, one per (year × type × split × topic)
├── build_inventory.py           regenerates inventory.csv from clef-tar/
├── clef-tar/                    full clone of CLEF-TAR (1.5 GB)
├── cohen-2006/epc-ir.clean.tsv  15-topic gold standard, 18,733 rows
├── run_validation.py            prep + score subcommands
├── orchestrate.py               drives the full E2E flow against a running app
├── build_slides.py              produces the validation slide deck
├── runs/                        output CSVs and JSON metric reports
├── figures/                     plots embedded in the slide deck
└── slides/                      PPTX presentation"""
    )

    doc.h2("run_validation.py prep")
    doc.p(
        "Given a CLEF-TAR topic ID, fetches every paper's title and abstract from PubMed "
        "and writes a CSV in ReviewPyper's expected column shape."
    )
    doc.code_with_explain(
        "    # validation/run_validation.py\n"
        "    parser.add_argument(\"topic\", help=\"CLEF-TAR topic ID, e.g. CD010355\")\n"
        "    parser.add_argument(\"--out\", required=True, help=\"Output CSV path\")\n"
        "    parser.add_argument(\"--email\", required=True, help=\"NCBI E-utilities requires an email\")\n"
        "    parser.add_argument(\"--year\", choices=[\"2017\", \"2018\", \"2019\"])\n"
        "    parser.add_argument(\"--split\", choices=[\"train\", \"test\"])\n"
        "    parser.add_argument(\"--max-pmids\", type=int)",
        "Topic ID is the only positional argument; year/split disambiguate when the same "
        "topic appears in multiple slices. --email is forwarded to NCBI per their fair-use "
        "policy. --max-pmids does a stratified sample (preserving include/exclude ratio) "
        "for fast experimentation.",
    )
    doc.h3("PubMed fetch loop")
    doc.code_with_explain(
        "EFETCH_BATCH = 200  # NCBI hard cap\n"
        "EFETCH_URL = \"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi\"\n\n"
        "def fetch_pubmed_records(pmids, email, ...):\n"
        "    for batch_idx in range(0, len(pmids), EFETCH_BATCH):\n"
        "        batch = pmids[batch_idx : batch_idx + EFETCH_BATCH]\n"
        "        params = urllib.parse.urlencode({\n"
        "            \"db\": \"pubmed\", \"id\": \",\".join(batch),\n"
        "            \"retmode\": \"xml\", \"rettype\": \"abstract\",\n"
        "            \"tool\": tool, \"email\": email,\n"
        "        })\n"
        "        body = _http_get_with_retry(EFETCH_URL + \"?\" + params)\n"
        "        out.update(_parse_efetch_xml(body))\n"
        "        time.sleep(PUBMED_DELAY_SEC)",
        "NCBI E-utilities accepts up to 200 PMIDs per request. We batch, parse XML, and "
        "sleep 0.4s between batches to stay safely under the 3-req/s rate limit. The XML "
        "parser handles structured AbstractText with section labels (Background:, Methods:, "
        "Results:, Conclusions:) and concatenates them into one Abstract field.",
    )

    doc.h2("orchestrate.py — full E2E against the running app")
    doc.p(
        "Talks to the API the same way the React frontend does: through nginx on "
        "localhost:3000/api. Three steps, plus a download to retrieve the screened CSV."
    )
    doc.code_with_explain(
        "def create_project(base, name):\n"
        "    body, ctype = _multipart_body({\"name\": name}, {})\n"
        "    raw = _http(\"POST\", f\"{base}/files/projects\", body, {\"Content-Type\": ctype})\n"
        "    return json.loads(raw)\n\n"
        "def upload_file(base, project_id, path):\n"
        "    body, ctype = _multipart_body(\n"
        "        {\"subfolder\": \"\"}, {\"file\": (path.name, path.read_bytes())}\n"
        "    )\n"
        "    raw = _http(\"POST\", f\"{base}/files/upload/{project_id}\", body,\n"
        "                {\"Content-Type\": ctype})\n"
        "    return json.loads(raw)\n\n"
        "def screen_titles(base, api_key_path, csv_path, question, model):\n"
        "    payload = {\n"
        "        \"api_key_path\": api_key_path, \"csv_path\": csv_path,\n"
        "        \"question\": question, \"model_choice\": model,\n"
        "    }\n"
        "    raw = _http(\"POST\", f\"{base}/titles/screen\",\n"
        "                json.dumps(payload).encode(),\n"
        "                {\"Content-Type\": \"application/json\"})\n"
        "    return json.loads(raw)",
        "Three thin wrappers around the corresponding routes. Notice the multipart body "
        "is built by hand (no requests dependency); we use stdlib only. The boundary is a "
        "secure-random hex token so no input could collide with it. screen_titles posts "
        "JSON because the pipeline routes accept Pydantic models, not form data — this is "
        "the same shape the frontend's pipelineApi.screenTitles sends.",
    )

    doc.h2("run_validation.py score")
    doc.p(
        "Compares the AI's predictions to the gold qrels. Computes recall, precision, F1, "
        "accuracy, and a binary approximation of WSS@95."
    )
    doc.code_with_explain(
        "def _read_predictions(path):\n"
        "    pred_col_candidates = (\"OpenAI_Screen\", \"AI_Screen\", \"Screen\", ...)\n"
        "    col = next(c for c in pred_col_candidates if c in reader.fieldnames)\n"
        "    for row in reader:\n"
        "        v = (row.get(col) or \"\").strip().lower()\n"
        "        if v in (\"1\", \"true\", \"yes\") or v.startswith(\"incl\"):\n"
        "            pred[pmid] = \"Include\"\n"
        "        elif v in (\"0\", \"false\", \"no\") or v.startswith(\"excl\"):\n"
        "            pred[pmid] = \"Exclude\"",
        "Tolerates several spellings of the prediction column and several spellings of "
        "the value (1/0, true/false, Include/Exclude). ReviewPyper writes 1/0 by default "
        "but earlier configs used Include/Exclude — handling both lets us re-run old "
        "outputs.",
    )
    doc.code_with_explain(
        "def _wss(labels, scores, target_recall=0.95):\n"
        "    pairs = sorted(zip(scores, labels), key=lambda p: -p[0])\n"
        "    tp = 0\n"
        "    for k, (_, lab) in enumerate(pairs, start=1):\n"
        "        if lab == 1:\n"
        "            tp += 1\n"
        "        if tp / n_pos >= target_recall:\n"
        "            return (n - k) / n - (1 - target_recall)\n"
        "    return 0.0",
        "Cohen 2006's Work Saved over Sampling at 95% recall: rank papers by predicted "
        "relevance, walk the ranking until you've recovered 95% of the true includes, and "
        "report the fraction of the corpus you saved over a random screen. Without "
        "continuous scores (ReviewPyper currently outputs binary 0/1), this collapses to "
        "a coarse approximation — but it's still informative as a sanity check.",
    )

    doc.h2("How to run a validation")
    doc.code(
        """# 1. Generate the input CSV from CLEF-TAR + PubMed
python3 validation/run_validation.py prep CD010355 \\
    --out validation/runs/CD010355_input.csv \\
    --email you@example.com --year 2019 --split train

# 2. Drive the full flow against the running app
python3 validation/orchestrate.py \\
    --input validation/runs/CD010355_input.csv \\
    --question "What is the effectiveness of NIPPV for prevention of...?" \\
    --api-base http://localhost:3000/api \\
    --out-screened validation/runs/CD010355_screened_gpt41.csv \\
    --model gpt4.1

# 3. Score against gold
python3 validation/run_validation.py score \\
    --predictions validation/runs/CD010355_screened_gpt41.csv \\
    --topic CD010355 --year 2019 --split train \\
    --report validation/runs/CD010355_metrics_gpt41.json

# 4. Build the slide deck
python3 validation/build_slides.py
"""
    )
    doc.callout(
        "Cost reference",
        "43-record run on gpt-3.5-turbo: ≈ $0.04. Same run on gpt-4.1: ≈ $0.70. "
        "A full Cochrane-scale topic (e.g. CD010705 with 228 records) on gpt-4.1: ≈ $4. "
        "$100 of OpenAI credit covers ~25 full-Cochrane-scale runs.",
    )

    doc.h2("First pilot results — CD010355")
    doc.p(
        "On topic CD010355 (NIPPV for post-pulmonary-resection complications, n=43, 9 "
        "gold-included), the off-the-shelf ReviewPyper title screener with two different "
        "models produced:"
    )
    doc.code(
        """Model            Recall   Precision   F1     Accuracy
gpt-3.5-turbo    0.444    0.333       0.381   0.698
gpt-4.1          0.667    0.750       0.706   0.884"""
    )
    doc.p(
        "GPT-4.1 nearly doubled F1 over GPT-3.5 with no other code change. For published "
        "Cochrane reviews, recall must reach ~0.95 — neither model crosses that bar yet, "
        "and prompt engineering plus active learning are obvious next levers. The full "
        "writeup lives in validation/slides/ReviewPyperAPI_validation.pptx."
    )
    doc.page_break()


def part5(doc: Doc):
    doc.h1("Part 5 — Defense  ·  questions to be ready for")
    doc.p(
        "Likely review-meeting questions, with the shape of the right answer."
    )

    questions = [
        ("Why two services (frontend + API) instead of one or three?",
         "One was untenable — the React app needs to be a single-page-app served as static files; mixing it into the same Python process means coupling build cycles. Three services (the previous gateway split-off) added latency and operational burden for no architectural gain. See ADR-0001."),
        ("What was the actual CSV-upload bug?",
         "services/api.ts set headers: { 'Content-Type': 'multipart/form-data' } explicitly. axios respects an explicit Content-Type even on FormData payloads, so the boundary parameter was missing and the body was unparseable. Removing the line lets axios set the correct boundary."),
        ("Why anyio.open_file instead of plain open() in upload_file?",
         "open() is synchronous. Inside an async route, sync I/O blocks the event loop, freezing the entire server for the duration of the write. anyio.open_file is non-blocking — other requests get serviced concurrently. For a 300 MB CSV that takes 15 seconds to write, that's 15 seconds of frozen API otherwise."),
        ("How does the test mode work end-to-end?",
         "useConfig.resolveEndpoint appends ?test=true to pipeline URLs when config.testMode is true (skipping /files/* and /health). On the server, TestModeMiddleware sees the query param, branches before the real route runs, writes synthetic CSVs / JSON to disk under /data/<project_id>/, and returns the same response shape the real route would produce. Result: full E2E flow without spending OpenAI tokens."),
        ("What's the source of truth for project state?",
         "Today: localStorage via useProjectState (legacy). Planned target: server-side state.json via useServerProjectState. The migration is intentionally incremental — see ADR-0002. The current dual-write means create_project always succeeds in both places."),
        ("Why semantic CSS tokens (--color-foreground) instead of raw scales (--color-gray-900)?",
         "A token re-naming for purpose vs. shade. When dark mode is added, only the @theme variable values need to change — every component using bg-[var(--color-surface)] adapts automatically. Components using bg-white wouldn't."),
        ("What does _safe_path protect against?",
         "Path traversal: a malicious or careless input like \"../../etc/passwd\" appended to DATA_DIR. After resolve(), the canonical path either is or isn't under DATA_DIR. is_relative_to checks correctly; the older str.startswith check has a false-positive on shared name prefixes (/data/foo vs /data/foobar)."),
        ("Why server-managed API keys (managed credits)?",
         "Onboarding friction. Asking a researcher to obtain an OpenAI key before they can run a single review kills adoption. With managed credits, the operator absorbs cost in exchange for usability. Power users keep the override path. See ADR-0005."),
        ("What does StrictMode do, and why?",
         "Development-only React feature that double-invokes effects and renders to surface bugs (use-after-unmount, derived-state effects, side effects in render). No-op in production. We wrap <App /> in it for development safety."),
        ("Why does useConfig throw when called outside a ConfigProvider?",
         "Loud failure beats silent misuse. If a developer accidentally renders a component using useConfig outside the provider tree (e.g. inside a Storybook story), they get an immediate error instead of an undefined config that crashes elsewhere later in a less obvious way."),
        ("What's the lazy-init pattern in useState and why does it matter?",
         "useState(() => expensiveCompute()) runs the function only on first render; useState(expensiveCompute()) runs it every render. When the initial state requires reading external mutable state (like localStorage), the lazy form is also the only correct option for render purity in React 19."),
        ("How is Tailwind v4 different from v3?",
         "v3 had a tailwind.config.ts file; v4 moves design tokens to a @theme block in CSS. v3 used @tailwind base/components/utilities; v4 uses @import \"tailwindcss\". v4 also drops the JIT engine entirely (it's the default and only engine). Our codebase is v4 native."),
        ("Why is the entire upload + screen flow split into two steps?",
         "Latency. Upload latency depends on file size; LLM-screening latency depends on row count and model. Combining them would mean a single user-facing button that takes 30s to 30min depending on the corpus. Splitting lets us show progress on each separately and lets the user pause / fix bad parameters between."),
    ]

    for q, a in questions:
        doc.h3("Q: " + q)
        doc.p("A. " + a)


# ─────────────────────── main ───────────────────────


def main():
    doc = Doc()
    title_page(doc)
    toc(doc)
    part0(doc)
    part1(doc)
    part2_intro(doc)
    part2_main_py(doc)
    part2_test_mode(doc)
    part2_files_router(doc)
    part3_intro(doc)
    part3_main_app(doc)
    part3_useconfig(doc)
    part3_api(doc)
    part3_useapi(doc)
    part3_useprojectstate(doc)
    part3_fileupload(doc)
    part3_stepper(doc)
    part3_setupreview(doc)
    part3_titlescreening(doc)
    part3_pdf(doc)
    part4(doc)
    part5(doc)
    part6_validation(doc)

    out = Path(__file__).resolve().parent / "ReviewPyPerAPI_line_by_line.docx"
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
